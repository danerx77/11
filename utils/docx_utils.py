"""
docx_utils.py – Narzędzia do generowania i wypełniania dokumentów DOCX.
Używa precyzyjnej modyfikacji na poziomie znaków,
aby zachować idealnie formatowanie i chronić tekst wokół tagów (np. (KW nr )).
"""

import re
import sys
import logging
import copy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- KONFIGURACJA LOGOWANIA BŁĘDÓW DO PLIKU ---
def setup_logger():
    if getattr(sys, 'frozen', False):
        base_dir = Path(sys.executable).parent.resolve()
    else:
        base_dir = Path(__file__).parent.parent.resolve()
    log_dir = base_dir / 'dane'
    log_dir.mkdir(parents=True, exist_ok=True)
    
    log_file = log_dir / 'error_log.txt'
    logging.basicConfig(
        filename=str(log_file),
        level=logging.ERROR,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        encoding='utf-8'
    )
    return logging.getLogger(__name__)

logger = setup_logger()

# ----------------------------------------------

def split_name_smartly(name: str, max_len: int) -> str:
    """
    Inteligentnie dzieli długą nazwę na dwie linie.
    Priorytety dzielenia:
    1. Spójnik " i " (najsilniejszy separator dla par/małżeństw)
    2. Przecinek "," (separator dla listy osób lub firmy i osoby)
    3. Spacja najbliżej środka tekstu (fallback)
    """
    if not name or len(name) <= max_len:
        return name

    # Priorytet 1: " i "
    if " i " in name:
        idx = name.find(" i ")
        return name[:idx].strip() + "\n" + name[idx:].strip()

    # Priorytet 2: Przecinek ","
    if "," in name:
        indices = [i for i, char in enumerate(name) if char == ',']
        middle = len(name) // 2
        # Wybieramy przecinek najbliżej środka, aby linie były w miarę równe
        best_comma = min(indices, key=lambda x: abs(x - middle))
        return name[:best_comma + 1].strip() + "\n" + name[best_comma + 1:].strip()

    # Priorytet 3: Spacja najbliżej środka (obecna logika)
    middle = len(name) // 2
    spaces = [i for i, char in enumerate(name) if char == ' ']
    if not spaces:
        return name
        
    best_space = min(spaces, key=lambda x: abs(x - middle))
    return name[:best_space].strip() + "\n" + name[best_space:].strip()

def format_long_name_for_envelope(name: str, max_len: int = 37) -> str:
    """
    Dzieli długie nazwy (np. pow. 37 znaków) znakiem nowej linii,
    który posłuży do utworzenia twardego entera (nowego akapitu) w Wordzie.
    """
    return split_name_smartly(name, max_len)

def format_long_name_for_cover(name: str, max_len: int = 34) -> str:
    """ 
    Formatuje długie nazwy adresata w Piśmie Przewodnim.
    Domyślny limit to 34 znaki w pierwszej linii.
    """
    return split_name_smartly(name, max_len)

def copy_run_format(src_run, dst_run):
    """Kopiuje formatowanie tekstu z jednego run-a do drugiego."""
    try:
        dst_run.bold = src_run.bold
        dst_run.italic = src_run.italic
        dst_run.underline = src_run.underline
        if src_run.font:
            if src_run.font.name:
                dst_run.font.name = src_run.font.name
            if src_run.font.size:
                dst_run.font.size = src_run.font.size
            if src_run.font.color and src_run.font.color.rgb:
                dst_run.font.color.rgb = src_run.font.color.rgb
    except Exception as e:
        logger.warning(f"Nie udało się skopiować formatowania czcionki: {e}")

def replace_text_precisely(elem, replacements: dict):
    for sdtPr in elem.findall(f'.//{qn("w:sdtPr")}'):
        plcHdr = sdtPr.find(f'.//{qn("w:showingPlcHdr")}')
        if plcHdr is not None: sdtPr.remove(plcHdr)
    wt_elems = list(elem.iter(qn('w:t')))
    if not wt_elems: return
    for key, value in replacements.items():
        if not key: continue
        val_str = str(value) if value is not None else ""
        
        while True:
            full_text = ""
            char_map = []
            for wt_idx, wt in enumerate(wt_elems):
                txt = wt.text or ""
                for char_idx in range(len(txt)):
                    full_text += txt[char_idx]
                    char_map.append((wt_idx, char_idx))
            
            if key not in full_text: break
                
            start_idx = full_text.find(key)
            end_idx = start_idx + len(key)
            wt_texts = [list(wt.text or "") for wt in wt_elems]
            
            for i in range(start_idx, end_idx):
                w_i, c_i = char_map[i]
                wt_texts[w_i][c_i] = ""
            
            w_start, c_start = char_map[start_idx]
            wt_texts[w_start][c_start] = val_str
            
            for wt_idx, wt in enumerate(wt_elems):
                new_text = "".join(wt_texts[wt_idx])
                if wt.text != new_text:
                    if "\n" in new_text:
                        parent_r = wt.getparent()
                        idx_in_parent = parent_r.index(wt)
                        parent_r.remove(wt)
                        
                        parts = new_text.split('\n')
                        for i, part in enumerate(parts):
                            if part or len(parts) == 1:
                                new_wt = OxmlElement('w:t')
                                new_wt.text = part
                                if part.startswith(" ") or part.endswith(" "):
                                    new_wt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                                parent_r.insert(idx_in_parent, new_wt)
                                idx_in_parent += 1
                                
                            if i < len(parts) - 1:
                                br = OxmlElement('w:br')
                                parent_r.insert(idx_in_parent, br)
                                idx_in_parent += 1
                    else:
                        wt.text = new_text
                        if new_text.startswith(" ") or new_text.endswith(" "):
                            wt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                            
            wt_elems = list(elem.iter(qn('w:t')))

def unlock_content_controls(doc: Document):
    settings_part = doc.settings.element
    for dp in settings_part.findall(qn('w:documentProtection')):
        settings_part.remove(dp)

def remove_paragraph_borders(doc: Document):
    for para in doc.paragraphs:
        pPr = para._p.find(qn('w:pPr'))
        if pPr is not None:
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is not None: pPr.remove(pBdr)

def _transform_comma_separated(value, transform):
    """Przetwarza każdą niepustą wartość z listy rozdzielonej przecinkami."""
    text = str(value or "").strip()
    if not text:
        return text

    parts = [part.strip() for part in text.split(",") if part.strip()]
    if not parts:
        return text
    return ", ".join(transform(part) for part in parts)


def apply_declension_preferences(
    location: str = "",
    street: str = "",
    county: str = "",
    preferences: dict = None,
):
    """Zwraca wartości pól po uwzględnieniu ustawień odmiany dokumentu.

    Przetwarzanie odbywa się tuż przed podmianą tagów DOCX. Dzięki temu te
    same ustawienia działają dla wszystkich standardowych tagów, niezależnie
    od sposobu generowania dokumentu (pojedynczo, seryjnie lub z własną mapą
    tagów).
    """
    location_value = str(location or "").strip()
    street_value = str(street or "").strip()
    county_value = str(county or "").strip()

    if not isinstance(preferences, dict):
        return location_value, street_value, county_value

    # Powiat wyznaczamy z pierwotnej miejscowości, a nie z jej już odmienionej
    # formy (np. Gdańsk -> gdański, a nie Gdańsku).
    source_location = location_value

    if preferences.get("decl_location_locative", False):
        from utils.polish_declension import decline_city

        location_value = _transform_comma_separated(location_value, decline_city)

    if preferences.get("decl_decline_streets", False):
        from utils.polish_declension import decline_street

        street_value = _transform_comma_separated(street_value, decline_street)

    if preferences.get("decl_powiat_zamiana", False):
        from utils.polish_declension import city_to_powiat

        county_source = county_value or source_location
        county_value = _transform_comma_separated(county_source, city_to_powiat)

    return location_value, street_value, county_value


# Akceptujemy oba warianty interpunkcji oraz starszy zapis bez polskich
# znaków. Dzięki temu standardowe pola działają także w istniejących szablonach.
STANDARD_TAG_ALIASES = {
    'location': (
        '<Miejscowość działki:>',
        '<Miejscowość działki>',
        '<Miejscowosc dzialki:>',
        '<Miejscowosc dzialki>',
    ),
    'address_street': ('<Ulica>', '<Ulica:>'),
    'street': ('<Ulica>', '<Ulica:>'),
    'county': ('<Powiat:>', '<Powiat>'),
}


def _build_replacements(default_tags: dict, values: dict, tag_map: dict = None) -> dict:
    """Tworzy podmiany dla standardowych i skonfigurowanych tagów.

    Własna mapa tagów nie usuwa obsługi domyślnego znacznika. Pozwala to
    bezpiecznie używać starszych szablonów, np. z <Miejscowość działki:>,
    po zmianie nazwy tagu w ustawieniach.
    """
    configured_tags = dict(default_tags)
    if isinstance(tag_map, dict):
        for key, new_tag in tag_map.items():
            if key not in configured_tags or not new_tag:
                continue
            configured_tags[key] = str(new_tag).strip()

    replacements = {}
    for key, value in values.items():
        tags = [default_tags.get(key), configured_tags.get(key)]
        tags.extend(STANDARD_TAG_ALIASES.get(key, ()))
        for tag in tags:
            if tag:
                replacements[str(tag)] = value
    return replacements


def generate_declaration(
    template_path: str, output_path: str, project_number: str = '', place: str = '', date_str: str = '', 
    owner_name: str = '', nip: str = '', pesel: str = '', location: str = '', 
    street: str = '', voivodeship: str = '', county: str = '', municipality: str = '', 
    parcel_numbers_budowa: str = '', parcel_numbers_demontaz: str = '', area_ha: str = '', 
    area_ha_budowa: str = '', area_ha_demontaz: str = '',
    precinct: str = '', precinct_caps: str = '', precinct_number: str = '', 
    kw_numbers: str = '', kw_numbers_budowa: str = '', kw_numbers_demontaz: str = '',
    device_description: str = '', declaration_type: str = 'budowa', tag_map: dict = None,
    unlock_docs: bool = False, declension_options: dict = None
) -> bool:
    try:
        doc = Document(template_path)
        location, street, county = apply_declension_preferences(
            location, street, county, declension_options
        )
        if not precinct_caps: precinct_caps = precinct.upper()
        default_tags = {
            'owner_name': '<Imie Nazwisko>', 'nip': '<Nip>', 'pesel': '<Pesel>',
            'voivodeship': '<Województwo:>', 'county': '<Powiat:>', 'municipality': '<Jednostka ewidencyjna:>',
            'location': '<Miejscowość działki:>', 'address_street': '<Ulica>', 
            'parcel_numbers_budowa': '<działki budowa:>', 'parcel_numbers_demontaz': '<działki demontaż:>',
            'area_ha': '<Powierzchnia [ha]>', 
            'area_ha_budowa': '<Powierzchnia Budowa>', 'area_ha_demontaz': '<Powierzchnia Demontaż>',
            'kw_numbers': '<<Nr KW>>', 'kw_numbers_budowa': '<<KW Budowa>>', 'kw_numbers_demontaz': '<<KW Demontaż>>',
            'device_description': '<wybór przykładów lub wpisać ręcznie>',
            'project_number': '<nr projektu>', 'date': '<Data>', 'place': '<Miejscowość>',
            'address': '<Adres>', 'precinct': '<Obręb ewidencyjny: wielka litery>', 'precinct_number': '<Nr obrębu>',
        }
        
        replacements = _build_replacements(
            default_tags,
            {
                'owner_name': owner_name,
                'nip': nip,
                'pesel': pesel,
                'voivodeship': voivodeship,
                'county': county,
                'municipality': municipality,
                'location': location,
                'address_street': street,
                'parcel_numbers_budowa': parcel_numbers_budowa,
                'parcel_numbers_demontaz': parcel_numbers_demontaz,
                'area_ha': area_ha,
                'area_ha_budowa': area_ha_budowa,
                'area_ha_demontaz': area_ha_demontaz,
                'kw_numbers': kw_numbers,
                'kw_numbers_budowa': kw_numbers_budowa,
                'kw_numbers_demontaz': kw_numbers_demontaz,
                'device_description': device_description,
                'project_number': project_number,
                'date': date_str,
                'place': place,
                'address': f"{street}, {location}" if street else location,
                'precinct': precinct,
                'precinct_number': precinct_number,
            },
            tag_map,
        )
        if declaration_type == 'demontaz':
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt
            target_tag = default_tags['owner_name']
            for p in doc.paragraphs:
                full_text = "".join([wt.text for wt in p._element.iter(qn('w:t')) if wt.text])
                if target_tag in full_text:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.left_indent = Pt(0)
                    p.paragraph_format.first_line_indent = Pt(0)
        replace_text_precisely(doc.element.body, replacements)
        for section in doc.sections:
            replace_text_precisely(section.header._element, replacements)
            replace_text_precisely(section.footer._element, replacements)
        
        if unlock_docs: unlock_content_controls(doc)
        doc.save(output_path)
        return True
    except Exception as e:
        logger.exception(f"Błąd podczas generowania Oświadczenia (generate_declaration): {e}")
        return False

def generate_cover_letter(
    template_path: str, output_path: str, date_str: str = '', place: str = '', addressee_salutation: str = 'Sz. P.', 
    addressee_name: str = '', addressee_street: str = '', addressee_city: str = '', location: str = '', 
    street: str = '', subject: str = '', task_construction: str = '', task_demolition: str = '',
    parcel_numbers_construction: list = None, parcel_numbers_demolition: list = None,
    parcel_numbers: list = None, ownership_phrase: str = 'których są Państwo współwłaścicielami,', 
    sender_name: str = '', sender_street: str = '', sender_city: str = '', tag_map: dict = None,
    unlock_docs: bool = False, declension_options: dict = None
) -> bool:
    try:
        doc = Document(template_path)
        location, street, _ = apply_declension_preferences(
            location, street, preferences=declension_options
        )
        unique_parcels = list(dict.fromkeys(parcel_numbers or []))
        all_nums_str = ', '.join(unique_parcels)
        parcel_type = 'działek nr' if len(unique_parcels) > 1 else 'działki nr'
        
        # Zastosowanie łamania długiej nazwy (limit 34 znaki w pierwszej linii)
        addressee_name = format_long_name_for_cover(addressee_name, max_len=34)
        default_tags = {
            'place': '<Miejscowość druku>', 'date': '<data sporządzenia>', 'sender_name': '<Imię Nazwisko nadawcy>',
            'sender_street': '<ul. nadawca>', 'sender_city': '<kod pocztowy Miejscowość nadawcy>',
            'addressee_name': '<Imię Nazwisko adresat>', 'addressee_street': '<ul. adresat>', 'addressee_city': '<kod pocztowy Miejscowość adresat>',
            'location': '<Miejscowość działki>', 'street': '<Ulica>', 'subject': '<Temat>',
            'task_construction': '<Zadanie budowa>', 'task_demolition': '<Zadanie demontaż>',
            'parcel_numbers_construction': '<działki budowa>', 'parcel_numbers_demolition': '<działki demontaż>',
            'parcel_type': '<odmiana działki>', 'parcel_numbers': '<Numer działki>',
            'ownership_phrase': '<wybór tekstu do kogo jest skierowany>'
        }
        
        replacements = _build_replacements(
            default_tags,
            {
                'place': place,
                'date': date_str,
                'sender_name': sender_name,
                'sender_street': sender_street,
                'sender_city': sender_city,
                'addressee_name': addressee_name,
                'addressee_street': addressee_street,
                'addressee_city': addressee_city,
                'location': location,
                'street': street,
                'subject': subject,
                'task_construction': task_construction,
                'task_demolition': task_demolition,
                'parcel_numbers_construction': ', '.join(parcel_numbers_construction or []),
                'parcel_numbers_demolition': ', '.join(parcel_numbers_demolition or []),
                'parcel_type': parcel_type,
                'parcel_numbers': all_nums_str,
                'ownership_phrase': ownership_phrase,
            },
            tag_map,
        )
        
        replace_text_precisely(doc.element.body, replacements)
        for section in doc.sections:
            replace_text_precisely(section.header._element, replacements)
            replace_text_precisely(section.footer._element, replacements)

        if unlock_docs: unlock_content_controls(doc)
        doc.save(output_path)
        return True
    except Exception as e: 
        logger.exception(f"Błąd podczas generowania Pisma Przewodniego (generate_cover_letter): {e}")
        return False

def generate_envelope_c5(template_path: str, output_path: str, sender_name: str = '', sender_street: str = '', sender_city: str = '', addressee_name: str = '', addressee_street: str = '', addressee_city: str = '', stamp_image_bytes: bytes = None) -> bool:
    try:
        doc = Document(template_path)
        
        paras = list(doc.paragraphs)
        p_sender_name = paras[0] if len(paras) > 0 else None
        p_sender_street = paras[1] if len(paras) > 1 else None
        p_sender_city = paras[2] if len(paras) > 2 else None
        
        p_addressee_name = paras[18] if len(paras) > 18 else None
        p_addressee_street = paras[19] if len(paras) > 19 else None
        p_addressee_city = paras[20] if len(paras) > 20 else None
        if stamp_image_bytes:
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    part = rel.target_part
                    part._blob = stamp_image_bytes
                    if hasattr(part, '_content_type') and ('emf' in (part._content_type or '').lower() or str(getattr(part, '_partname', '')).lower().endswith('.emf')):
                        part._content_type = 'image/png'
                        old_pn = str(getattr(part, '_partname', ''))
                        if old_pn.lower().endswith('.emf'):
                            from docx.opc.packuri import PackURI
                            part._partname = PackURI(old_pn[:-4] + '.png')
                    break
                    
        if p_sender_name and len(p_sender_name.runs) >= 3: p_sender_name.runs[2].text = sender_name
        if p_sender_street and p_sender_street.runs: p_sender_street.runs[0].text = sender_street; [setattr(r, 'text', '') for r in p_sender_street.runs[1:]]
        if p_sender_city and p_sender_city.runs: p_sender_city.runs[0].text = sender_city; [setattr(r, 'text', '') for r in p_sender_city.runs[1:]]
        
        if p_addressee_name:
            formatted_name = format_long_name_for_envelope(addressee_name, max_len=37)
            lines = formatted_name.split('\n')
            
            if len(lines) > 1:
                p_addressee_name.runs[0].text = lines[0]
                for r in p_addressee_name.runs[1:]: r.text = ''
                
                new_p_element = copy.deepcopy(p_addressee_name._element)
                for r in new_p_element.findall(qn('w:r')):
                    new_p_element.remove(r)
                
                p_addressee_name._element.addnext(new_p_element)
                
                from docx.text.paragraph import Paragraph
                new_p = Paragraph(new_p_element, p_addressee_name._parent)
                new_run = new_p.add_run(lines[1])
                
                if p_addressee_name.runs:
                    copy_run_format(p_addressee_name.runs[0], new_run)
            else:
                p_addressee_name.runs[0].text = formatted_name
                for r in p_addressee_name.runs[1:]: r.text = ''
        
        if p_addressee_street and p_addressee_street.runs: p_addressee_street.runs[0].text = addressee_street; [setattr(r, 'text', '') for r in p_addressee_street.runs[1:]]
        if p_addressee_city and p_addressee_city.runs: p_addressee_city.runs[0].text = addressee_city; [setattr(r, 'text', '') for r in p_addressee_city.runs[1:]]
        
        remove_paragraph_borders(doc)
        doc.save(output_path)
        return True
    except Exception as e:
        logger.exception(f"Błąd podczas generowania Koperty C5: {e}")
        return False

def generate_envelope_c6(template_path: str, output_path: str, addressee_name: str = '', addressee_street: str = '', addressee_city: str = '', stamp_image_bytes: bytes = None) -> bool:
    try:
        doc = Document(template_path)
        
        paras = list(doc.paragraphs)
        p_addressee_name = paras[6] if len(paras) > 6 else None
        p_addressee_street = paras[8] if len(paras) > 8 else None
        p_addressee_city = paras[9] if len(paras) > 9 else None
        
        if stamp_image_bytes:
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    part = rel.target_part
                    part._blob = stamp_image_bytes
                    if hasattr(part, '_content_type') and ('emf' in (part._content_type or '').lower() or str(getattr(part, '_partname', '')).lower().endswith('.emf')):
                        part._content_type = 'image/png'
                        old_pn = str(getattr(part, '_partname', ''))
                        if old_pn.lower().endswith('.emf'):
                            from docx.opc.packuri import PackURI
                            part._partname = PackURI(old_pn[:-4] + '.png')
                    break
                    
        if p_addressee_name:
            formatted_name = format_long_name_for_envelope(addressee_name, max_len=37)
            lines = formatted_name.split('\n')
            
            if len(lines) > 1 and len(p_addressee_name.runs) >= 5:
                p_addressee_name.runs[4].text = lines[0]
                
                new_p_element = copy.deepcopy(p_addressee_name._element)
                for r in new_p_element.findall(qn('w:r')):
                    new_p_element.remove(r)
                
                p_addressee_name._element.addnext(new_p_element)
                
                from docx.text.paragraph import Paragraph
                new_p = Paragraph(new_p_element, p_addressee_name._parent)
                new_run = new_p.add_run(lines[1])
                
                copy_run_format(p_addressee_name.runs[4], new_run)
            else:
                if len(p_addressee_name.runs) >= 5:
                    p_addressee_name.runs[4].text = formatted_name
        
        if p_addressee_street and len(p_addressee_street.runs) >= 4: p_addressee_street.runs[3].text = f'\t{addressee_street}'
        if p_addressee_city and len(p_addressee_city.runs) >= 5: p_addressee_city.runs[4].text = addressee_city
        
        doc.save(output_path)
        return True
    except Exception as e:
        logger.exception(f"Błąd podczas generowania Koperty C6: {e}")
        return False

def batch_change_date(docx_paths: list, old_date: str, new_date: str) -> list:
    modified = []
    for path in docx_paths:
        try:
            doc = Document(path)
            found = False
            for para in doc.paragraphs:
                full = ''.join(r.text for r in para.runs)
                if old_date in full:
                    new_full = full.replace(old_date, new_date)
                    if para.runs:
                        para.runs[0].text = new_full
                        for r in para.runs[1:]: r.text = ''
                    found = True
            if found:
                doc.save(path)
                modified.append(path)
        except Exception as e:
            logger.exception(f"Błąd zmiany daty w pliku {path}: {e}")
    return modified

def merge_docx_files(file_paths: list, output_path: str) -> bool:
    if not file_paths: return False
    try:
        import win32com.client
        import os
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        master = word.Documents.Open(os.path.abspath(file_paths[0]), ReadOnly=True)
        master.SaveAs(os.path.abspath(output_path))
        master.Close()
        master_doc = word.Documents.Open(os.path.abspath(output_path))
        selection = word.Selection
        for file in file_paths[1:]:
            file_abs = os.path.abspath(file)
            if not os.path.exists(file_abs): continue
            selection.EndKey(Unit=6)
            selection.InsertBreak(Type=7)
            selection.InsertFile(FileName=file_abs)
        master_doc.Save()
        master_doc.Close()
        word.Quit()
        return True
    except Exception as e: 
        logger.exception(f"Błąd łączenia plików przez COM (Word): {e}")
        try:
            from docx import Document
            from docxcompose.composer import Composer
            import os
            master = Document(os.path.abspath(file_paths[0]))
            composer = Composer(master)
            for file in file_paths[1:]:
                file_abs = os.path.abspath(file)
                if not os.path.exists(file_abs): continue
                doc_to_append = Document(file_abs)
                master.add_page_break()
                composer.append(doc_to_append)
            master.save(os.path.abspath(output_path))
            return True
        except Exception as e2:
            logger.exception(f"Błąd łączenia plików przez docxcompose: {e2}")
            return False

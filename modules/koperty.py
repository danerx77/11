import json
import os
import re
import shutil
import sys
import uuid
from datetime import datetime
from pathlib import Path

from utils.global_settings import (
    get_global_data_dir,
    load_global_envelope_preferences,
    load_global_stamp_settings,
    save_global_envelope_preferences,
    save_global_stamp_settings,
)
from PySide6.QtCore import QByteArray, QSize, Qt, QThread, Signal
from PySide6.QtGui import QColor, QIcon, QKeySequence, QPixmap, QShortcut
from PySide6.QtWidgets import (
    QAbstractItemView,
    QCheckBox,
    QComboBox,
    QDialog,
    QDialogButtonBox,
    QFileDialog,
    QFormLayout,
    QGroupBox,
    QHBoxLayout,
    QHeaderView,
    QLabel,
    QLineEdit,
    QListView,
    QListWidget,
    QListWidgetItem,
    QMessageBox,
    QProgressDialog,
    QPushButton,
    QScrollArea,
    QSizePolicy,
    QSpinBox,
    QSplitter,
    QTabWidget,
    QTableWidget,
    QTableWidgetItem,
    QVBoxLayout,
    QWidget,
)

class StampLoaderWorker(QThread):
    finished = Signal(list)
    error = Signal(str)

    def __init__(
        self,
        pdf_path: str,
        envelope_type: str,
        stamp_profile: dict | None = None,
    ):
        super().__init__()
        self.pdf_path = pdf_path
        self.envelope_type = envelope_type
        self.stamp_profile = stamp_profile

    def run(self):
        try:
            from utils.pdf_utils import extract_stamps_from_pdf, render_stamp_thumbnail
            stamps = extract_stamps_from_pdf(
                self.pdf_path,
                self.envelope_type,
                self.stamp_profile,
            )
            for stamp in stamps:
                stamp["thumbnail_bytes"] = render_stamp_thumbnail(
                    stamp["pdf_path"],
                    stamp["page"],
                    stamp["rect"],
                    size=150,
                )
            self.finished.emit(stamps)
        except Exception as error:
            self.error.emit(str(error))

class StampVisualSettingsDialog(QDialog):
    def __init__(self, parent, pdf_path: str, env_type: str, config: dict):
        super().__init__(parent)
        self.pdf_path = pdf_path
        self.env_type = env_type
        self.config = config
        self.profile_key = f"stamp_profile_{env_type.lower()}"
        self.setWindowTitle(f"Ustawienia wycinania dla {env_type}")
        self.resize(850, 500)
        geometry = self.config.get(f"stamp_window_geom_{env_type}", "")
        if geometry:
            try:
                self.restoreGeometry(QByteArray.fromHex(geometry.encode()))
            except Exception:
                pass

        main_layout = QHBoxLayout(self)
        options_layout = QFormLayout()
        spin_style = (
            "QSpinBox { padding-right: 25px; padding-left: 5px; "
            "min-height: 25px; }"
        )

        if env_type == "C5":
            default_profile = {
                "crop_left": 90,
                "crop_right": 90,
                "crop_up": 136,
                "crop_down": 2,
            }
            info = QLabel(
                "TRYB C5: ROZCIĄGANIE "
                "(odniesienie: dół kodu kreskowego)<br>"
                "Po kliknięciu OK profil jest zapisywany globalnie w "
                "dane/stamp_profiles.json."
            )
            ranges = [(0, 300), (0, 300), (-50, 300), (-50, 300)]
        else:
            default_profile = {
                "crop_left": 0,
                "crop_right": 0,
                "crop_up": 0,
                "crop_down": 0,
            }
            info = QLabel(
                "TRYB C6: DOCINANIE KRAWĘDZI OBRAZKA<br>"
                "Po kliknięciu OK profil jest zapisywany globalnie w "
                "dane/stamp_profiles.json."
            )
            ranges = [(-100, 100)] * 4

        info.setStyleSheet("font-weight: bold; color: #a4b0be;")
        info.setWordWrap(True)
        options_layout.addRow(info)

        self.profile = self.config.get(self.profile_key, default_profile)
        self.spin_left = QSpinBox()
        self.spin_right = QSpinBox()
        self.spin_up = QSpinBox()
        self.spin_down = QSpinBox()
        spins = [
            self.spin_left,
            self.spin_right,
            self.spin_up,
            self.spin_down,
        ]
        keys = ["crop_left", "crop_right", "crop_up", "crop_down"]
        labels = [
            "Krawędź LEWA (px):",
            "Krawędź PRAWA (px):",
            "Krawędź GÓRNA (px):",
            "Krawędź DOLNA (px):",
        ]
        for spin, key, label, limits in zip(spins, keys, labels, ranges):
            spin.setRange(*limits)
            spin.setStyleSheet(spin_style)
            spin.setValue(self.profile.get(key, default_profile[key]))
            spin.valueChanged.connect(self._update_preview)
            options_layout.addRow(label, spin)

        left_widget = QWidget()
        left_widget.setLayout(options_layout)
        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok
            | QDialogButtonBox.StandardButton.Cancel
        )
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)

        left_layout = QVBoxLayout()
        left_layout.addWidget(left_widget)
        left_layout.addStretch()
        left_layout.addWidget(buttons)

        self.lbl_preview = QLabel("Ładowanie podglądu...")
        self.lbl_preview.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.lbl_preview.setStyleSheet(
            "background-color: white; border: 2px solid #555; color: black;"
        )
        self.lbl_preview.setMinimumSize(350, 350)

        main_layout.addLayout(left_layout, 1)
        main_layout.addWidget(self.lbl_preview, 2)
        self._update_preview()

    def _update_preview(self):
        if not self.pdf_path or not Path(self.pdf_path).exists():
            self.lbl_preview.setText("Brak PDF do podglądu.")
            return
        from utils.pdf_utils import render_stamp_preview
        image_bytes = render_stamp_preview(
            self.pdf_path,
            self.env_type,
            self.get_profile(),
        )
        if image_bytes:
            pixmap = QPixmap()
            pixmap.loadFromData(image_bytes)
            self.lbl_preview.setPixmap(pixmap)
        else:
            self.lbl_preview.setText(
                "Nie wycięto znaczka.\nZbyt ciasne marginesy?"
            )

    def get_profile(self) -> dict:
        return {
            "crop_left": self.spin_left.value(),
            "crop_right": self.spin_right.value(),
            "crop_up": self.spin_up.value(),
            "crop_down": self.spin_down.value(),
        }

    def _save_geometry(self):
        self.config[f"stamp_window_geom_{self.env_type}"] = (
            self.saveGeometry().toHex().data().decode()
        )

    def accept(self):
        self._save_geometry()
        self.config[self.profile_key] = self.get_profile()
        super().accept()

    def reject(self):
        self._save_geometry()
        super().reject()

class EnvelopeGenWidget(QWidget):
    shipment_generated = Signal(dict)

    def __init__(self, config: dict, parent=None):
        super().__init__(parent)
        self.config = config
        self.owners: list[dict] = []
        self.stamps_c5: list[dict] = []
        self.stamps_c6: list[dict] = []
        self.active_project_path = None
        self._worker_c5 = None
        self._worker_c6 = None
        self._progress_c5 = None
        self._progress_c6 = None
        self.setAcceptDrops(True)
        self._ensure_global_data_dir()
        self._load_global_stamp_profiles()
        self._load_global_envelope_preferences()
        self._build_ui()

    # ------------------------------------------------------------------ UI
    def _build_ui(self):
        main_layout = QHBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        self.main_splitter = QSplitter(Qt.Orientation.Horizontal)
        self.main_splitter.setObjectName("envelope_main_splitter")
        self.main_splitter.setChildrenCollapsible(False)
        self.main_splitter.setHandleWidth(7)
        main_layout.addWidget(self.main_splitter)

        # LEWY PANEL – PRZEWIJANY
        left_widget = QWidget()
        left_widget.setObjectName("envelope_left_content")
        left_widget.setMinimumWidth(520)
        left_widget.setSizePolicy(
            QSizePolicy.Policy.Expanding,
            QSizePolicy.Policy.Preferred,
        )
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(8, 8, 8, 8)
        left_layout.setSpacing(10)

        header = QLabel("✉️ Inteligentny Generator Kopert C5 / C6")
        header.setStyleSheet("font-size: 16px; font-weight: 700;")
        left_layout.addWidget(header)

        owner_box = self._build_owners_box()
        owner_box.setSizePolicy(
            QSizePolicy.Policy.Expanding,
            QSizePolicy.Policy.Expanding,
        )
        left_layout.addWidget(owner_box, 1)

        left_layout.addWidget(self._build_sender_box())
        left_layout.addWidget(self._build_templates_box())
        left_layout.addWidget(self._build_generation_box())
        left_layout.addStretch()

        self.left_scroll = QScrollArea()
        self.left_scroll.setObjectName("envelope_left_scroll")
        self.left_scroll.setWidgetResizable(True)
        self.left_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        self.left_scroll.setHorizontalScrollBarPolicy(
            Qt.ScrollBarPolicy.ScrollBarAsNeeded
        )
        self.left_scroll.setVerticalScrollBarPolicy(
            Qt.ScrollBarPolicy.ScrollBarAsNeeded
        )
        self.left_scroll.setMinimumWidth(520)
        self.left_scroll.setWidget(left_widget)
        self.main_splitter.addWidget(self.left_scroll)

        # PRAWY PANEL
        right_widget = QWidget()
        right_widget.setObjectName("envelope_right_content")
        right_widget.setMinimumWidth(420)
        right_widget.setSizePolicy(
            QSizePolicy.Policy.Expanding,
            QSizePolicy.Policy.Expanding,
        )
        right_layout = QVBoxLayout(right_widget)
        right_layout.setContentsMargins(6, 8, 8, 8)
        right_layout.addWidget(self._build_stamps_box())
        self.main_splitter.addWidget(right_widget)

        self.main_splitter.setStretchFactor(0, 3)
        self.main_splitter.setStretchFactor(1, 2)
        saved_sizes = self.config.get("envelope_splitter_sizes", [])
        if (
            isinstance(saved_sizes, list)
            and len(saved_sizes) == 2
            and all(isinstance(size, int) and size > 0 for size in saved_sizes)
        ):
            self.main_splitter.setSizes(saved_sizes)
        else:
            self.main_splitter.setSizes([700, 500])
        self.main_splitter.splitterMoved.connect(self._save_splitter_sizes)
        self._auto_load_stamps_from_ia(silent=True)

    def _build_owners_box(self) -> QGroupBox:
        box = QGroupBox("Lista Adresatów (zapisywana w projekcie)")
        layout = QVBoxLayout(box)
        layout.setSpacing(10)
        layout.setContentsMargins(10, 15, 10, 10)

        load_row = QHBoxLayout()
        load_external = QPushButton("Wczytaj TXT/XLSX")
        load_external.clicked.connect(self._load_external_list)
        load_folder = QPushButton("Z folderu Pism DOCX")
        load_folder.setStyleSheet(
            "background-color: #2b5797; color: white; font-weight: bold;"
        )
        load_folder.clicked.connect(self._load_from_docx_folder)
        load_files = QPushButton("Z plików Pism DOCX")
        load_files.setStyleSheet(
            "background-color: #2b5797; color: white; font-weight: bold;"
        )
        load_files.clicked.connect(self._load_from_docx_files)
        load_row.addWidget(load_external)
        load_row.addWidget(load_folder)
        load_row.addWidget(load_files)
        layout.addLayout(load_row)

        filter_row = QHBoxLayout()
        self.chk_hide_generated = QCheckBox(
            "Ukryj adresatów, dla których wygenerowano obie koperty"
        )
        self.chk_hide_generated.setChecked(
            bool(self.config.get("envelope_hide_generated", False))
        )
        self.chk_hide_generated.toggled.connect(self._on_owner_filter_changed)
        filter_row.addWidget(self.chk_hide_generated)
        self.chk_show_only_generated = QCheckBox("Pokaż tylko wygenerowane")
        self.chk_show_only_generated.setChecked(
            bool(self.config.get("envelope_show_only_generated", False))
        )
        self.chk_show_only_generated.toggled.connect(self._on_owner_filter_changed)
        filter_row.addWidget(self.chk_show_only_generated)
        filter_row.addStretch()
        filter_row.addWidget(QLabel("Sortowanie:"))
        self.view_sort_combo = QComboBox()
        self.view_sort_combo.addItems(["Domyślne", "Alfabetycznie"])
        self.view_sort_combo.setCurrentIndex(
            self._preference_index("envelope_view_sort")
        )
        self.view_sort_combo.currentIndexChanged.connect(
            self._on_view_sort_changed
        )
        self.view_sort_combo.setToolTip(
            "Wybrane sortowanie jest zapisywane globalnie w dane/envelope_preferences.json."
        )
        filter_row.addWidget(self.view_sort_combo)
        remembered_note = QLabel("↻ zapamiętywane globalnie")
        remembered_note.setToolTip(
            "Sortowanie, filtry widoku, tryb generowania i układ kolumn będą "
            "takie same po ponownym uruchomieniu oraz w następnym projekcie."
        )
        remembered_note.setStyleSheet("color:#78909c; font-size:11px;")
        filter_row.addWidget(remembered_note)
        layout.addLayout(filter_row)

        self.table_owners = QTableWidget(0, 4)
        self.table_owners.setHorizontalHeaderLabels(
            ["Koperta C5", "Koperta C6", "Adresat", "Adres"]
        )
        header = self.table_owners.horizontalHeader()
        header.setSectionsMovable(True)
        saved_table_state = self.config.get("envelope_table_state", "")
        if saved_table_state:
            try:
                header.restoreState(QByteArray.fromHex(saved_table_state.encode()))
            except Exception:
                pass
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        header.setSectionResizeMode(2, QHeaderView.ResizeMode.Stretch)
        header.setSectionResizeMode(3, QHeaderView.ResizeMode.Stretch)
        header.sectionMoved.connect(self._save_owners_table_state)
        header.sectionResized.connect(self._save_owners_table_state)

        self.table_owners.setMinimumHeight(250)
        self.table_owners.setSizePolicy(
            QSizePolicy.Policy.Expanding,
            QSizePolicy.Policy.Expanding,
        )
        self.table_owners.setAlternatingRowColors(True)
        self.table_owners.setSelectionBehavior(
            QAbstractItemView.SelectionBehavior.SelectRows
        )
        self.table_owners.setSelectionMode(
            QAbstractItemView.SelectionMode.ExtendedSelection
        )
        self.table_owners.setEditTriggers(
            QAbstractItemView.EditTrigger.NoEditTriggers
        )
        self.table_owners.itemSelectionChanged.connect(
            self._on_table_selection
        )
        QShortcut(QKeySequence("Delete"), self.table_owners).activated.connect(
            self._delete_selected_owners
        )
        layout.addWidget(self.table_owners, 1)
        layout.addSpacing(10)

        status_row = QHBoxLayout()
        for text, env_type, state in [
            ("✅ C5: TAK", "C5", True),
            ("❌ C5: NIE", "C5", False),
            ("✅ C6: TAK", "C6", True),
            ("❌ C6: NIE", "C6", False),
        ]:
            button = QPushButton(text)
            button.clicked.connect(
                lambda checked=False, e=env_type, s=state:
                self._set_manual_status(e, s)
            )
            status_row.addWidget(button)
        layout.addLayout(status_row)

        select_row = QHBoxLayout()
        select_all = QPushButton("☑️ Zaznacz widocznych")
        select_all.clicked.connect(self.table_owners.selectAll)
        select_missing = QPushButton("🔍 Zaznacz bez koperty")
        select_missing.clicked.connect(self._select_missing_envelopes)
        delete = QPushButton("🗑️ Usuń zaznaczonych")
        delete.setObjectName("btn_danger")
        delete.clicked.connect(self._delete_selected_owners)
        select_row.addWidget(select_all)
        select_row.addWidget(select_missing)
        select_row.addWidget(delete)
        layout.addLayout(select_row)

        layout.addSpacing(10)
        form_container = QWidget()
        form = QFormLayout(form_container)
        form.setContentsMargins(0, 0, 0, 0)
        self.addr_name_edit = QLineEdit()
        self.addr_street_edit = QLineEdit()
        self.addr_city_edit = QLineEdit()
        form.addRow("Nazwa adresata:", self.addr_name_edit)
        form.addRow("Ulica:", self.addr_street_edit)
        form.addRow("Kod i miejscowość:", self.addr_city_edit)
        layout.addWidget(form_container)

        return box

    def _build_sender_box(self) -> QGroupBox:
        box = QGroupBox("Dane Nadawcy (na C5 i powrót C6)")
        form = QFormLayout(box)
        sender = self._sender_preference()
        self.sender_name_edit = QLineEdit(sender.get("name", ""))
        self.sender_street_edit = QLineEdit(sender.get("street", ""))
        self.sender_city_edit = QLineEdit(sender.get("city", ""))
        for edit in (
            self.sender_name_edit,
            self.sender_street_edit,
            self.sender_city_edit,
        ):
            edit.editingFinished.connect(self._save_sender_preferences)
        form.addRow("Imię i nazwisko:", self.sender_name_edit)
        form.addRow("Ulica:", self.sender_street_edit)
        form.addRow("Kod i miejscowość:", self.sender_city_edit)
        return box

    def _path_row(self, line_edit, browse, reset, clear):
        widget = QWidget()
        layout = QHBoxLayout(widget)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(line_edit, 1)
        for icon, tip, callback in [
            ("📂", "Wybierz plik", browse),
            ("↩", "Przywróć ustawienie domyślne", reset),
            ("✖", "Usuń ścieżkę", clear),
        ]:
            button = QPushButton(icon)
            button.setToolTip(tip)
            button.setFixedWidth(42)
            button.clicked.connect(callback)
            layout.addWidget(button)
        return widget

    def _build_templates_box(self) -> QGroupBox:
        box = QGroupBox("Szablony kopert (.docx)")
        form = QFormLayout(box)
        self.c5_tmpl_edit = QLineEdit(
            self._template_preference("C5")
        )
        self.c6_tmpl_edit = QLineEdit(
            self._template_preference("C6")
        )
        self.c5_tmpl_edit.textChanged.connect(
            lambda value: self._save_template_preference("C5", value)
        )
        self.c6_tmpl_edit.textChanged.connect(
            lambda value: self._save_template_preference("C6", value)
        )
        form.addRow(
            "Szablon C5:",
            self._path_row(
                self.c5_tmpl_edit,
                lambda: self._browse_template("C5"),
                self._set_default_templates,
                self.c5_tmpl_edit.clear,
            ),
        )
        form.addRow(
            "Szablon C6:",
            self._path_row(
                self.c6_tmpl_edit,
                lambda: self._browse_template("C6"),
                self._set_default_templates,
                self.c6_tmpl_edit.clear,
            ),
        )
        return box

    def _build_generation_box(self) -> QGroupBox:
        box = QGroupBox("Akcje generowania")
        layout = QVBoxLayout(box)
        output_row = QHBoxLayout()
        output_row.addWidget(QLabel("Folder zapisu:"))
        self.out_dir_edit = QLineEdit(
            self.config.get("envelope_output_dir", "")
        )
        self.out_dir_edit.setReadOnly(True)
        self.out_dir_edit.textChanged.connect(
            lambda value: self.config.update(
                {"envelope_output_dir": value.strip()}
            )
        )
        browse = QPushButton("Wybierz")
        browse.clicked.connect(self._browse_output_dir)
        clear = QPushButton("Usuń")
        clear.clicked.connect(self.out_dir_edit.clear)
        output_row.addWidget(self.out_dir_edit, 1)
        output_row.addWidget(browse)
        output_row.addWidget(clear)
        layout.addLayout(output_row)

        options = QHBoxLayout()
        self.chk_single_files = QCheckBox(
            "Każda koperta w osobnym pliku"
        )
        self.chk_single_files.setChecked(
            bool(self.config.get("envelope_single_files", False))
        )
        self.chk_single_files.toggled.connect(self._on_single_files_toggled)
        options.addWidget(self.chk_single_files)
        options.addStretch()
        options.addWidget(QLabel("Kolejność:"))
        self.sort_combo = QComboBox()
        self.sort_combo.addItems(["Jak na liście", "Alfabetycznie"])
        self.sort_combo.setCurrentIndex(
            self._preference_index("envelope_generation_sort")
        )
        self.sort_combo.currentIndexChanged.connect(
            self._on_generation_sort_changed
        )
        self.sort_combo.setToolTip(
            "Kolejność generowania jest zapamiętywana globalnie."
        )
        options.addWidget(self.sort_combo)
        layout.addLayout(options)

        single_type_row = QHBoxLayout()
        self.btn_gen_c5 = QPushButton(
            "⚡ Generuj tylko C5 dla zaznaczonych"
        )
        self.btn_gen_c5.setObjectName("btn_primary")
        self.btn_gen_c5.clicked.connect(self._generate_c5_envelopes)
        single_type_row.addWidget(self.btn_gen_c5)
        self.btn_gen_c6 = QPushButton(
            "⚡ Generuj tylko C6 dla zaznaczonych"
        )
        self.btn_gen_c6.setObjectName("btn_primary")
        self.btn_gen_c6.clicked.connect(self._generate_c6_envelopes)
        single_type_row.addWidget(self.btn_gen_c6)
        layout.addLayout(single_type_row)

        self.btn_gen_batch = QPushButton(
            "⚡ Generuj C5 + C6 dla zaznaczonych adresatów"
        )
        self.btn_gen_batch.setObjectName("btn_accent")
        self.btn_gen_batch.setMinimumHeight(38)
        self.btn_gen_batch.clicked.connect(self._generate_batch_envelopes)
        layout.addWidget(self.btn_gen_batch)

        self.btn_regenerate = QPushButton(
            "🔄 Wygeneruj ponownie, używając tych samych znaczków"
        )
        self.btn_regenerate.setStyleSheet(
            "background-color: #f39c12; color: white; font-weight: bold;"
        )
        self.btn_regenerate.clicked.connect(self._regenerate_envelopes)
        layout.addWidget(self.btn_regenerate)
        return box

    def _build_stamps_box(self) -> QGroupBox:
        box = QGroupBox("Zarządzanie Znaczkami (pamięć globalna)")
        layout = QVBoxLayout(box)
        form = QFormLayout()
        self.c5_pdf_edit = QLineEdit(self.config.get("stamp_c5_pdf", ""))
        self.c5_pdf_edit.setReadOnly(True)
        form.addRow("Znaczki C5:", self._stamp_pdf_row("C5"))
        self.c6_pdf_edit = QLineEdit(self.config.get("stamp_c6_pdf", ""))
        self.c6_pdf_edit.setReadOnly(True)
        form.addRow("Znaczki C6:", self._stamp_pdf_row("C6"))
        layout.addLayout(form)

        self.stamps_tab = QTabWidget()
        self.tab_c5, self.list_c5 = self._stamp_list_page()
        self.tab_c6, self.list_c6 = self._stamp_list_page()
        self.stamps_tab.addTab(self.tab_c5, "Znaczki C5")
        self.stamps_tab.addTab(self.tab_c6, "Znaczki C6")
        self.stamps_tab.setCurrentIndex(
            self._preference_index("envelope_stamps_tab", default=0)
        )
        self.stamps_tab.currentChanged.connect(self._on_stamps_tab_changed)
        layout.addWidget(self.stamps_tab, 1)

        buttons = QHBoxLayout()
        used = QPushButton("✅ Użyte")
        used.clicked.connect(self._mark_selected_used)
        unused = QPushButton("⭕ Wolne")
        unused.clicked.connect(self._mark_selected_unused)
        reset = QPushButton("🗑️ Resetuj")
        reset.clicked.connect(self._reset_stamp_usage)
        settings = QPushButton("⚙️ Ustawienia cięcia")
        settings.clicked.connect(self._open_visual_settings)
        for button in (used, unused, reset, settings):
            buttons.addWidget(button)
        layout.addLayout(buttons)

        auto_load = QPushButton("🔍 Wymuś załadowanie znaczków")
        auto_load.setObjectName("btn_primary")
        auto_load.clicked.connect(self._auto_load_stamps_from_ia)
        layout.addWidget(auto_load)

        self.lbl_stamp_info = QLabel(
            "Wczytaj pliki PDF ze znaczkami, aby zarządzać kodami."
        )
        self.lbl_stamp_info.setWordWrap(True)
        self.lbl_stamp_info.setStyleSheet("color: #aaa; font-size: 11px;")
        layout.addWidget(self.lbl_stamp_info)
        return box

    def _stamp_pdf_row(self, env_type: str):
        widget = QWidget()
        layout = QHBoxLayout(widget)
        layout.setContentsMargins(0, 0, 0, 0)
        edit = self.c5_pdf_edit if env_type == "C5" else self.c6_pdf_edit
        layout.addWidget(edit, 1)
        choose = QPushButton("📂")
        choose.setToolTip("Wybierz PDF")
        choose.clicked.connect(lambda: self._browse_stamp_pdf(env_type))
        refresh = QPushButton("⟳")
        refresh.setToolTip("Odśwież znaczki")
        refresh.clicked.connect(lambda: self._load_stamps(edit.text(), env_type))
        clear = QPushButton("✖")
        clear.setToolTip("Usuń PDF z widoku")
        clear.clicked.connect(lambda: self._clear_stamps(env_type))
        for button in (choose, refresh, clear):
            button.setFixedWidth(42)
            layout.addWidget(button)
        return widget

    def _stamp_list_page(self):
        page = QWidget()
        layout = QVBoxLayout(page)
        stamp_list = QListWidget()
        stamp_list.setViewMode(QListView.ViewMode.IconMode)
        stamp_list.setResizeMode(QListView.ResizeMode.Adjust)
        stamp_list.setIconSize(QSize(130, 130))
        stamp_list.setGridSize(QSize(150, 180))
        stamp_list.setSelectionMode(
            QAbstractItemView.SelectionMode.ExtendedSelection
        )
        layout.addWidget(stamp_list)
        return page, stamp_list

    # ----------------------------------------------------------- lifecycle
    def _load_global_envelope_preferences(self):
        """Ładuje globalne sortowanie i wybory widoku przed zbudowaniem UI."""
        preferences = load_global_envelope_preferences(self._get_global_data_dir())
        if preferences:
            self.config.update(preferences)

    def _save_global_envelope_preferences(self) -> bool:
        """Trwale zapamiętuje wybory Kopert niezależnie od projektu."""
        return save_global_envelope_preferences(
            self.config,
            self._get_global_data_dir(),
        )

    def _preference_index(self, key: str, default=None) -> int:
        if default is None:
            default = 1 if self.config.get("sort_alpha_default", False) else 0
        try:
            index = int(self.config.get(key, default))
        except (TypeError, ValueError):
            index = default
        return 1 if index else 0

    def _on_owner_filter_changed(self, _checked=False):
        self.config["envelope_hide_generated"] = self.chk_hide_generated.isChecked()
        self.config["envelope_show_only_generated"] = (
            self.chk_show_only_generated.isChecked()
        )
        self._save_global_envelope_preferences()
        self._refresh_owners_table()

    def _on_view_sort_changed(self, index: int):
        self.config["envelope_view_sort"] = index
        self._save_global_envelope_preferences()
        self._refresh_owners_table()

    def _on_generation_sort_changed(self, index: int):
        self.config["envelope_generation_sort"] = index
        self._save_global_envelope_preferences()

    def _on_single_files_toggled(self, checked: bool):
        self.config["envelope_single_files"] = bool(checked)
        self._save_global_envelope_preferences()

    def _on_stamps_tab_changed(self, index: int):
        self.config["envelope_stamps_tab"] = int(index)
        self._save_global_envelope_preferences()

    def _save_owners_table_state(self, *_args):
        header = self.table_owners.horizontalHeader()
        self.config["envelope_table_state"] = (
            header.saveState().toHex().data().decode()
        )
        self._save_global_envelope_preferences()

    def _save_splitter_sizes(self, *_args):
        # Układ paneli jest wygodnym ustawieniem interfejsu, lecz pozostaje
        # lokalny dla sesji; globalnie zapisywane są wybory sortowania/widoku.
        self.config["envelope_splitter_sizes"] = self.main_splitter.sizes()

    def _template_preference(self, env_type: str) -> str:
        module_value = str(
            self.config.get(f"envelope_{env_type.lower()}_template", "") or ""
        ).strip()
        return module_value or self.config.get(
            f"env_{env_type.lower()}_template", ""
        )

    def _save_template_preference(self, env_type: str, path: str):
        self.config[f"envelope_{env_type.lower()}_template"] = path.strip()

    def _sender_preference(self) -> dict:
        saved_sender = self.config.get("envelope_sender")
        if isinstance(saved_sender, dict):
            return saved_sender
        return self.config.get("sender", {})

    def _save_sender_preferences(self):
        # Dane wpisane w Kopertach są lokalnym, zapamiętanym ustawieniem tego
        # modułu. Nie nadpisują formularza nadawcy dla innych dokumentów.
        sender = dict(self._sender_preference())
        sender.update(
            {
                "name": self.sender_name_edit.text().strip(),
                "street": self.sender_street_edit.text().strip(),
                "city": self.sender_city_edit.text().strip(),
            }
        )
        self.config["envelope_sender"] = sender

    def showEvent(self, event):
        super().showEvent(event)
        for combo, key in (
            (self.view_sort_combo, "envelope_view_sort"),
            (self.sort_combo, "envelope_generation_sort"),
        ):
            combo.blockSignals(True)
            combo.setCurrentIndex(self._preference_index(key))
            combo.blockSignals(False)

        for checkbox, key in (
            (self.chk_hide_generated, "envelope_hide_generated"),
            (self.chk_show_only_generated, "envelope_show_only_generated"),
            (self.chk_single_files, "envelope_single_files"),
        ):
            checkbox.blockSignals(True)
            checkbox.setChecked(bool(self.config.get(key, False)))
            checkbox.blockSignals(False)

        sender = self._sender_preference()
        for edit, key in (
            (self.sender_name_edit, "name"),
            (self.sender_street_edit, "street"),
            (self.sender_city_edit, "city"),
        ):
            value = sender.get(key, "")
            if not edit.hasFocus() and edit.text() != value:
                edit.setText(value)

        c5_template = self._template_preference("C5")
        c6_template = self._template_preference("C6")
        if self.c5_tmpl_edit.text() != c5_template:
            self.c5_tmpl_edit.setText(c5_template)
        if self.c6_tmpl_edit.text() != c6_template:
            self.c6_tmpl_edit.setText(c6_template)
        output_dir = self.config.get("envelope_output_dir", "")
        if self.out_dir_edit.text() != output_dir:
            self.out_dir_edit.setText(output_dir)
        self._sync_pdf_from_config("C5")
        self._sync_pdf_from_config("C6")

    def _sync_pdf_from_config(self, env_type):
        key = f"stamp_{env_type.lower()}_pdf"
        path = self.config.get(key, "")
        edit = self.c5_pdf_edit if env_type == "C5" else self.c6_pdf_edit
        if path != edit.text():
            edit.setText(path)
            if path and Path(path).exists():
                self._load_stamps(path, env_type)

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()

    def dropEvent(self, event):
        for url in event.mimeData().urls():
            if not url.isLocalFile():
                continue
            path = url.toLocalFile()
            suffix = Path(path).suffix.lower()
            lower = path.lower()
            if suffix == ".docx":
                if "c5" in lower or "wysył" in lower:
                    self.c5_tmpl_edit.setText(path)
                elif "c6" in lower or "zwrot" in lower:
                    self.c6_tmpl_edit.setText(path)
            elif suffix == ".pdf":
                env_type = "C6" if "c6" in lower else "C5"
                edit = self.c6_pdf_edit if env_type == "C6" else self.c5_pdf_edit
                edit.setText(path)
                self.config[f"stamp_{env_type.lower()}_pdf"] = path
                self._load_stamps(path, env_type)
            elif suffix in {".txt", ".csv", ".xlsx"}:
                self._load_external_list_from_path(path)
            event.acceptProposedAction()
            return

    # -------------------------------------------------------------- owners
    def _get_unique_selected_rows(self) -> list[int]:
        model = self.table_owners.selectionModel()
        indexes = model.selectedRows() or model.selectedIndexes()
        return sorted({index.row() for index in indexes})

    def _set_manual_status(self, env_type: str, state: bool):
        rows = self._get_unique_selected_rows()
        if not rows:
            QMessageBox.warning(self, "Brak zaznaczenia", "Zaznacz adresatów.")
            return
        for row in rows:
            item = self.table_owners.item(row, 2)
            data = item.data(Qt.ItemDataRole.UserRole) if item else None
            if data:
                index, _address = data
                self.owners[index][f"env_{env_type.lower()}_generated"] = state
                self.owners[index][f"manual_env_{env_type.lower()}"] = True
        self._save_owners_state()
        self._refresh_owners_table()

    def _delete_selected_owners(self):
        rows = self._get_unique_selected_rows()
        if not rows:
            QMessageBox.warning(self, "Brak", "Wybierz adresatów do usunięcia.")
            return
        if QMessageBox.question(
            self,
            "Usuwanie",
            f"Czy usunąć {len(rows)} zaznaczonych pozycji?",
        ) != QMessageBox.StandardButton.Yes:
            return
        indices = set()
        for row in rows:
            item = self.table_owners.item(row, 2)
            data = item.data(Qt.ItemDataRole.UserRole) if item else None
            if data:
                indices.add(data[0])
        for index in sorted(indices, reverse=True):
            if 0 <= index < len(self.owners):
                del self.owners[index]
        self._save_owners_state()
        self._refresh_owners_table()

    def _save_owners_state(self):
        if not self.active_project_path:
            return
        try:
            path = Path(self.active_project_path) / "adresaci.json"
            path.parent.mkdir(parents=True, exist_ok=True)
            path.write_text(
                json.dumps(self.owners, ensure_ascii=False, indent=4),
                encoding="utf-8",
            )
        except Exception:
            pass

    def _load_owners_state(self):
        self.owners = []
        if self.active_project_path:
            path = Path(self.active_project_path) / "adresaci.json"
            if path.exists():
                try:
                    data = json.loads(path.read_text(encoding="utf-8"))
                    if isinstance(data, list):
                        self.owners = data
                except Exception:
                    pass
        self._refresh_owners_table()

    def _add_unique_owners(self, new_owners: list) -> int:
        added = 0
        existing = {
            (
                self._get_formatted_name(owner).strip().casefold(),
                owner.get("address", "").strip().casefold(),
            )
            for owner in self.owners
        }
        for owner in new_owners:
            if not isinstance(owner, dict):
                continue
            key = (
                self._get_formatted_name(owner).strip().casefold(),
                owner.get("address", "").strip().casefold(),
            )
            if key[0] and key not in existing:
                self.owners.append(owner)
                existing.add(key)
                added += 1
        self._save_owners_state()
        self._refresh_owners_table()
        return added

    def _get_formatted_name(self, owner: dict) -> str:
        name = owner.get("full_name", "").strip()
        if owner.get("first_name") or owner.get("last_name"):
            try:
                from utils.gender_utils import format_salutation_line
                people = [owner]
                if owner.get("partner"):
                    people.append(owner["partner"])
                lines = format_salutation_line(people).split("\n")
                formatted = lines[1] if len(lines) >= 2 else lines[0]
                if formatted.strip():
                    name = formatted.strip()
            except Exception:
                pass
        return name

    def _get_short_name(self, owner: dict) -> str:
        first = owner.get("first_name", "").strip()
        last = owner.get("last_name", "").strip()
        if not first and not last:
            return owner.get("full_name", "Nieznany").replace(" ", "_")[:40]
        if " i " in first.lower() or owner.get("is_couple"):
            initials = [word[0].upper() for word in first.split() if word.lower() != "i"]
            return f"{'.'.join(initials)}.{last}"
        return f"{first[0].upper() if first else ''}.{last}"

    def _load_shipments(self) -> list:
        if not self.active_project_path:
            return []
        path = Path(self.active_project_path) / "przesylki" / "shipments.json"
        try:
            return json.loads(path.read_text(encoding="utf-8")) if path.exists() else []
        except Exception:
            return []

    def _refresh_owners_table(self):
        self.table_owners.setRowCount(0)
        shipments = self._load_shipments()
        display = []
        for index, owner in enumerate(self.owners):
            name = self._get_formatted_name(owner)
            addresses = [owner.get("address", "")]
            if owner.get("address_2"):
                addresses.append(owner["address_2"])
            for address in addresses:
                display.append((index, owner, name, address))

        if self.view_sort_combo.currentIndex() == 1:
            display.sort(key=lambda value: self._get_short_name(value[1]).casefold())

        for index, owner, name, address in display:
            has_c5 = owner.get("env_c5_generated", False)
            has_c6 = owner.get("env_c6_generated", False)
            if not owner.get("manual_env_c5"):
                has_c5 = has_c5 or any(
                    item.get("addressee") == name
                    and item.get("envelope_type") == "C5"
                    for item in shipments
                )
            if not owner.get("manual_env_c6"):
                has_c6 = has_c6 or any(
                    name in item.get("addressee", "")
                    and item.get("envelope_type") == "C6"
                    for item in shipments
                )
            owner["env_c5_generated"] = has_c5
            owner["env_c6_generated"] = has_c6
            if self.chk_hide_generated.isChecked() and has_c5 and has_c6:
                continue
            if hasattr(self, 'chk_show_only_generated') and self.chk_show_only_generated.isChecked() and not (has_c5 or has_c6):
                continue
            flags = []
            color = QColor("#2ecc71")
            if owner.get("is_dead"):
                flags.append("[ZMARŁY/A]")
                color = QColor("#e74c3c")
            elif owner.get("is_institution"):
                flags.append("[INSTYTUCJA]")
                color = QColor("#9b5de5")
            elif owner.get("is_church"):
                flags.append("[PARAFIA]")
                color = QColor("#f39c12")
            elif owner.get("is_company"):
                flags.append("[FIRMA]")
                color = QColor("#3498db")
            elif not address.strip():
                flags.append("[BRAK ADRESU]")
                color = QColor("#e67e22")
            elif not re.search(r"\d{2}-\d{3}", address):
                flags.append("[BŁĄD KODU]")
                color = QColor("#e67e22")

            row = self.table_owners.rowCount()
            self.table_owners.insertRow(row)
            c5_item = QTableWidgetItem("✅ TAK" if has_c5 else "⭕ NIE")
            c6_item = QTableWidgetItem("✅ TAK" if has_c6 else "⭕ NIE")
            c5_item.setForeground(QColor("#1dd1a1" if has_c5 else "#ff6b6b"))
            c6_item.setForeground(QColor("#1dd1a1" if has_c6 else "#ff6b6b"))
            name_item = QTableWidgetItem(f"{' '.join(flags)} {name}".strip())
            name_item.setData(Qt.ItemDataRole.UserRole, (index, address))
            name_item.setForeground(color)
            self.table_owners.setItem(row, 0, c5_item)
            self.table_owners.setItem(row, 1, c6_item)
            self.table_owners.setItem(row, 2, name_item)
            self.table_owners.setItem(row, 3, QTableWidgetItem(address))

    def _select_missing_envelopes(self):
        from PySide6.QtCore import QItemSelection, QItemSelectionModel
        # Czyścimy obecne zaznaczenie
        self.table_owners.clearSelection()
        
        # Tworzymy zbiorczy obiekt zaznaczenia
        selection = QItemSelection()
        model = self.table_owners.model()
        
        for row in range(self.table_owners.rowCount()):
            c5 = self.table_owners.item(row, 0)
            c6 = self.table_owners.item(row, 1)
            
            # Sprawdzamy, czy w którejś z kolumn statusu jest "NIE"
            if (c5 and "NIE" in c5.text()) or (c6 and "NIE" in c6.text()):
                # Dodajemy zakres od pierwszej do ostatniej kolumny danego wiersza
                left_index = model.index(row, 0)
                right_index = model.index(row, self.table_owners.columnCount() - 1)
                selection.select(left_index, right_index)
        
        # Aplikujemy całe zaznaczenie za jednym razem korzystając z poprawnego enuma
        self.table_owners.selectionModel().select(
            selection, 
            QItemSelectionModel.SelectionFlag.Select | QItemSelectionModel.SelectionFlag.Rows
        )

    def _on_table_selection(self):
        rows = self._get_unique_selected_rows()
        if not rows:
            return
        item = self.table_owners.item(rows[0], 2)
        data = item.data(Qt.ItemDataRole.UserRole) if item else None
        if not data:
            return
        index, address = data
        self.addr_name_edit.setText(self._get_formatted_name(self.owners[index]))
        if "," in address:
            street, city = address.rsplit(",", 1)
            self.addr_street_edit.setText(street.strip())
            self.addr_city_edit.setText(city.strip())
        else:
            self.addr_street_edit.clear()
            self.addr_city_edit.setText(address.strip())

    # --------------------------------------------------------- import lists
    def _load_from_docx_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Wybierz folder DOCX")
        if folder:
            files = [p for p in Path(folder).glob("*.docx") if not p.name.startswith("~$")]
            self._process_docx_files(self._natural_sort(files))

    def _load_from_docx_files(self):
        files, _ = QFileDialog.getOpenFileNames(
            self, "Wybierz pisma", "", "Word (*.docx)"
        )
        self._process_docx_files(self._natural_sort([Path(p) for p in files]))

    @staticmethod
    def _natural_sort(paths):
        return sorted(
            paths,
            key=lambda p: [
                int(part) if part.isdigit() else part.casefold()
                for part in re.split(r"([0-9]+)", p.name)
            ],
        )

    def _process_docx_files(self, files: list[Path]):
        if not files:
            QMessageBox.information(self, "Brak plików", "Nie znaleziono DOCX.")
            return
        try:
            from docx import Document
        except ImportError:
            QMessageBox.critical(self, "Błąd", "Zainstaluj python-docx.")
            return
        progress = QProgressDialog("Analizowanie pism...", "Anuluj", 0, len(files), self)
        progress.setWindowModality(Qt.WindowModality.WindowModal)
        owners = []
        for number, path in enumerate(files):
            if progress.wasCanceled():
                break
            progress.setValue(number)
            try:
                document = Document(path)
                lines = []
                for paragraph in document.paragraphs:
                    lines.extend(x.strip() for x in paragraph.text.splitlines() if x.strip())
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                lines.extend(x.strip() for x in paragraph.text.splitlines() if x.strip())
                owner = self._extract_address_from_lines(lines)
                if owner:
                    owners.append(owner)
            except Exception as error:
                print(f"Błąd analizy {path.name}: {error}")
        progress.setValue(len(files))
        added = self._add_unique_owners(owners)
        QMessageBox.information(self, "Import", f"Dodano {added} nowych adresatów.")

    @staticmethod
    def _extract_address_from_lines(lines: list[str]):
        for index, line in enumerate(lines):
            if not re.match(r"^Sz\.?\s*P\.?(?:\s|$)", line, re.I):
                continue
            block = []
            first = re.sub(r"(?i)^Sz\.?\s*P\.?\s*", "", line).strip()
            if first:
                block.append(first)
            for following in lines[index + 1:index + 8]:
                if re.match(r"^(Dotyczy|Znak|Działając|W związku|Nasz znak)", following, re.I):
                    break
                block.append(following)
            postal = next((i for i, text in enumerate(block) if re.search(r"\d{2}-\d{3}", text)), -1)
            if postal >= 0:
                city = block[postal]
                before = block[:postal]
                street = before[-1] if len(before) > 1 and re.search(r"\d", before[-1]) else ""
                name_parts = before[:-1] if street else before
                name = " ".join(name_parts).strip()
                if name:
                    return {"full_name": name, "address": f"{street}, {city}".strip(", ")}
            elif block:
                return {"full_name": block[0], "address": ", ".join(block[1:])}
        return None

    def _load_external_list(self):
        path, _ = QFileDialog.getOpenFileName(
            self,
            "Wybierz listę adresatów",
            "",
            "Listy (*.txt *.csv* .xlsx)",
        )
        if path:
            self._load_external_list_from_path(path)

    def _load_external_list_from_path(self, path: str):
        owners = []
        try:
            suffix = Path(path).suffix.lower()
            if suffix in {".txt", ".csv"}:
                with open(path, "r", encoding="utf-8", errors="replace") as file:
                    for line in file:
                        parts = [part.strip() for part in line.split(",")]
                        if parts and parts[0]:
                            owners.append({
                                "full_name": parts[0],
                                "address": ", ".join(parts[1:]).strip(", "),
                            })
            elif suffix == ".xlsx":
                import openpyxl
                workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
                for row in workbook.active.iter_rows(values_only=True):
                    if row and row[0]:
                        owners.append({
                            "full_name": str(row[0]),
                            "address": ", ".join(
                                str(value) for value in row[1:3] if value
                            ),
                        })
            added = self._add_unique_owners(owners)
            QMessageBox.information(self, "Import", f"Dodano {added} adresatów.")
        except Exception as error:
            QMessageBox.critical(self, "Błąd", f"Nie udało się wczytać pliku:\n{error}")

    # -------------------------------------------------------------- stamps
    def _get_global_data_dir(self) -> Path:
        return get_global_data_dir()

    def _ensure_global_data_dir(self):
        try:
            self._get_global_data_dir().mkdir(parents=True, exist_ok=True)
        except OSError:
            pass

    def _load_global_stamp_profiles(self):
        saved_profiles = load_global_stamp_settings(self._get_global_data_dir())
        if saved_profiles:
            self.config.update(saved_profiles)

    def _save_global_stamp_profiles(self) -> bool:
        return save_global_stamp_settings(self.config, self._get_global_data_dir())

    def _get_stamps_global_file(self) -> Path:
        return self._get_global_data_dir() / "stamps.json"

    def _clear_stamps(self, env_type):
        if env_type == "C5":
            self.c5_pdf_edit.clear()
            self.stamps_c5.clear()
            self.list_c5.clear()
        else:
            self.c6_pdf_edit.clear()
            self.stamps_c6.clear()
            self.list_c6.clear()
        self.config[f"stamp_{env_type.lower()}_pdf"] = ""
        self._update_stamps_info()

    def _browse_stamp_pdf(self, env_type):
        from utils.templates import (
            STAMP_FOLDER_NAMES,
            resolve_template_start_directory,
        )

        edit = self.c5_pdf_edit if env_type == "C5" else self.c6_pdf_edit
        start_dir = resolve_template_start_directory(
            self.config,
            config_key="path_znaczki",
            folder_names=STAMP_FOLDER_NAMES,
            current_path=edit.text(),
        )
        path, _ = QFileDialog.getOpenFileName(
            self,
            f"Wybierz PDF {env_type}",
            str(start_dir),
            "PDF (*.pdf)",
        )
        if path:
            edit.setText(path)
            self.config[f"stamp_{env_type.lower()}_pdf"] = path
            self._load_stamps(path, env_type)

    def _load_stamps(self, pdf_path: str, env_type: str):
        if not pdf_path or not Path(pdf_path).exists():
            return
        worker_name = f"_worker_{env_type.lower()}"
        progress_name = f"_progress_{env_type.lower()}"
        previous = getattr(self, worker_name, None)
        if previous and previous.isRunning():
            previous.requestInterruption()
            previous.quit()
            previous.wait(1000)
        progress = QProgressDialog(f"Wczytywanie znaczków {env_type}...", None, 0, 0, self)
        progress.setWindowModality(Qt.WindowModality.WindowModal)
        progress.show()
        setattr(self, progress_name, progress)
        worker = StampLoaderWorker(
            pdf_path,
            env_type,
            self.config.get(f"stamp_profile_{env_type.lower()}"),
        )
        worker.finished.connect(lambda stamps, kind=env_type: self._on_stamps_loaded(stamps, kind))
        worker.error.connect(self._on_stamps_error)
        setattr(self, worker_name, worker)
        worker.start()

    def _on_stamps_loaded(self, stamps: list, env_type: str):
        progress = getattr(self, f"_progress_{env_type.lower()}", None)
        if progress:
            progress.close()
        saved = self._load_saved_stamp_status(env_type)
        for index, stamp in enumerate(stamps):
            key = stamp.get("barcode") or f"idx_{index}"
            if key in saved:
                stamp["used"] = saved[key].get("used", False)
                stamp["used_by"] = saved[key].get("used_by", "")
        if env_type == "C5":
            self.stamps_c5 = stamps
        else:
            self.stamps_c6 = stamps
        self._refresh_stamps_list(env_type)
        self._update_stamps_info()
        self._save_stamp_status()

    def _on_stamps_error(self, message: str):
        for name in ("_progress_c5", "_progress_c6"):
            progress = getattr(self, name, None)
            if progress:
                progress.close()
        QMessageBox.critical(self, "Błąd", f"Nie udało się wczytać znaczków:\n{message}")

    def _refresh_stamps_list(self, env_type: str):
        widget = self.list_c5 if env_type == "C5" else self.list_c6
        stamps = self.stamps_c5 if env_type == "C5" else self.stamps_c6
        widget.clear()
        for index, stamp in enumerate(stamps):
            item = QListWidgetItem()
            item.setData(Qt.ItemDataRole.UserRole, stamp)
            pixmap = QPixmap()
            if stamp.get("thumbnail_bytes"):
                pixmap.loadFromData(stamp["thumbnail_bytes"])
            barcode = stamp.get("barcode", "").replace("(00)", "")
            barcode = ("..." + barcode[-10:]) if len(barcode) > 10 else (barcode or "[brak kodu]")
            item.setIcon(QIcon(pixmap))
            item.setText(f"#{index + 1}\n{barcode}\n[{'UŻYTY' if stamp.get('used') else 'WOLNY'}]")
            item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            item.setBackground(QColor("#5c2020" if stamp.get("used") else "#1e3d2f"))
            item.setForeground(QColor("#ff8888" if stamp.get("used") else "#88ff88"))
            widget.addItem(item)

    def _update_stamps_info(self):
        c5_free = sum(not stamp.get("used") for stamp in self.stamps_c5)
        c6_free = sum(not stamp.get("used") for stamp in self.stamps_c6)
        self.lbl_stamp_info.setText(
            f"C5: {len(self.stamps_c5)} (wolnych {c5_free}) | "
            f"C6: {len(self.stamps_c6)} (wolnych {c6_free})"
        )

    def _active_stamp_objects(self):
        if self.stamps_tab.currentIndex() == 0:
            return "C5", self.list_c5, self.stamps_c5
        return "C6", self.list_c6, self.stamps_c6

    def _mark_selected_used(self):
        env_type, widget, stamps = self._active_stamp_objects()
        for item in widget.selectedItems():
            index = widget.row(item)
            stamps[index]["used"] = True
            stamps[index]["used_by"] = "Ręcznie oznaczony"
        self._refresh_stamps_list(env_type)
        self._update_stamps_info()
        self._save_stamp_status()

    def _mark_selected_unused(self):
        env_type, widget, stamps = self._active_stamp_objects()
        for item in widget.selectedItems():
            index = widget.row(item)
            stamps[index]["used"] = False
            stamps[index]["used_by"] = ""
        self._refresh_stamps_list(env_type)
        self._update_stamps_info()
        self._save_stamp_status()

    def _reset_stamp_usage(self):
        env_type, widget, stamps = self._active_stamp_objects()
        if not stamps or QMessageBox.question(self, "Reset", f"Zresetować {env_type}?") != QMessageBox.StandardButton.Yes:
            return
        for stamp in stamps:
            stamp["used"] = False
            stamp["used_by"] = ""
        self._refresh_stamps_list(env_type)
        self._update_stamps_info()
        self._save_stamp_status()

    def _save_stamp_status(self):
        path = self._get_stamps_global_file()
        try:
            existing = json.loads(path.read_text(encoding="utf-8")) if path.exists() else {}
            for env_type, stamps, pdf in [
                ("C5", self.stamps_c5, self.c5_pdf_edit.text()),
                ("C6", self.stamps_c6, self.c6_pdf_edit.text()),
            ]:
                if not stamps or not pdf:
                    continue
                section = existing.setdefault(pdf, {})
                for index, stamp in enumerate(stamps):
                    key = stamp.get("barcode") or f"idx_{index}"
                    section[key] = {
                        "used": stamp.get("used", False),
                        "used_by": stamp.get("used_by", ""),
                        "env_type": env_type,
                    }
            path.parent.mkdir(parents=True, exist_ok=True)
            path.write_text(json.dumps(existing, ensure_ascii=False, indent=4), encoding="utf-8")
        except Exception:
            pass

    def _load_saved_stamp_status(self, env_type: str) -> dict:
        path = self._get_stamps_global_file()
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
            pdf = self.c5_pdf_edit.text() if env_type == "C5" else self.c6_pdf_edit.text()
            return data.get(pdf, {})
        except Exception:
            return {}

    def _get_unused_stamp(self, env_type: str):
        stamps = self.stamps_c5 if env_type == "C5" else self.stamps_c6
        return next((stamp for stamp in stamps if not stamp.get("used")), None)

    def _open_visual_settings(self):
        env_type = "C5" if self.stamps_tab.currentIndex() == 0 else "C6"
        pdf = self.c5_pdf_edit.text() if env_type == "C5" else self.c6_pdf_edit.text()
        if not pdf or not Path(pdf).exists():
            QMessageBox.warning(self, "Brak pliku", f"Wczytaj PDF {env_type}.")
            return
        dialog = StampVisualSettingsDialog(self, pdf, env_type, self.config)
        accepted = dialog.exec() == QDialog.DialogCode.Accepted
        if accepted:
            # Dialog zapisuje profil do współdzielonej konfiguracji. Zapisz go
            # od razu także do dane/stamp_profiles.json, niezależnie od projektu.
            if not self._save_global_stamp_profiles():
                QMessageBox.warning(
                    self,
                    "Nie zapisano profilu",
                    "Nie udało się zapisać ustawień cięcia w dane/stamp_profiles.json.",
                )
            self._load_stamps(pdf, env_type)

    def _auto_load_stamps_from_ia(self, silent=False):
        from utils.templates import (
            STAMP_FOLDER_NAMES,
            resolve_template_start_directory,
        )

        for env_type, edit in [("C5", self.c5_pdf_edit), ("C6", self.c6_pdf_edit)]:
            if edit.text() and Path(edit.text()).exists():
                stamps = self.stamps_c5 if env_type == "C5" else self.stamps_c6
                if not stamps:
                    self._load_stamps(edit.text(), env_type)

        directory = resolve_template_start_directory(
            self.config,
            config_key="path_znaczki",
            folder_names=STAMP_FOLDER_NAMES,
        )
        if not directory.is_dir():
            return

        pdfs = [
            path for path in directory.iterdir()
            if path.is_file() and path.suffix.lower() == ".pdf"
        ]
        for env_type, edit, words in [
            ("C5", self.c5_pdf_edit, ("c5", "neoznacze")),
            ("C6", self.c6_pdf_edit, ("c6",)),
        ]:
            if edit.text():
                continue
            found = next(
                (path for path in pdfs if any(word in path.name.lower() for word in words)),
                None,
            )
            if found:
                edit.setText(str(found))
                self.config[f"stamp_{env_type.lower()}_pdf"] = str(found)
                self._load_stamps(str(found), env_type)

    # ------------------------------------------------------------ generate
    def _browse_template(self, env_type: str):
        from utils.templates import (
            EXAMPLES_FOLDER_NAMES,
            resolve_template_start_directory,
        )

        edit = self.c5_tmpl_edit if env_type == "C5" else self.c6_tmpl_edit
        start_dir = resolve_template_start_directory(
            self.config,
            config_key="path_przyklady",
            folder_names=EXAMPLES_FOLDER_NAMES,
            current_path=edit.text(),
        )
        path, _ = QFileDialog.getOpenFileName(
            self,
            f"Szablon {env_type}",
            str(start_dir),
            "Word (*.docx)",
        )
        if path:
            edit.setText(path)

    def _set_default_templates(self):
        self.c5_tmpl_edit.setText(self._template_preference("C5"))
        self.c6_tmpl_edit.setText(self._template_preference("C6"))

    def _browse_output_dir(self):
        folder = QFileDialog.getExistingDirectory(self, "Wybierz folder zapisu")
        if folder:
            self.out_dir_edit.setText(folder)

    def _get_output_dir(self) -> Path:
        custom = self.out_dir_edit.text().strip()
        if custom:
            path = Path(custom)
        elif self.active_project_path:
            path = Path(self.active_project_path) / "koperty"
        else:
            path = Path.cwd()
        path.mkdir(parents=True, exist_ok=True)
        return path

    @staticmethod
    def _split_address(address):
        if "," in address:
            street, city = address.rsplit(",", 1)
            return street.strip(), city.strip()
        return "", address.strip()

    @staticmethod
    def _safe_filename(value):
        return re.sub(r'[\\/*?:"<>|\r\n]+', "", value).strip()

    def _selected_targets(self, validate=False):
        result = []
        for row in self._get_unique_selected_rows():
            item = self.table_owners.item(row, 2)
            data = item.data(Qt.ItemDataRole.UserRole) if item else None
            if data:
                index, address = data
                if not validate or self._is_valid_for_gen(self.owners[index], address)[0]:
                    result.append((index, address))
        if self.sort_combo.currentIndex() == 1:
            result.sort(key=lambda value: self._get_short_name(self.owners[value[0]]).casefold())
        return result

    def _is_valid_for_gen(self, owner: dict, address: str):
        if owner.get("is_dead"):
            return False, "Osoba zmarła"
        if not address.strip():
            return False, "Brak adresu"
        if not re.search(r"\d{2}-\d{3}", address):
            return False, "Brak kodu pocztowego"
        return True, "OK"

    def _generation_inputs(self, envelope_types=("C5", "C6")):
        """Waliduje wyłącznie szablony potrzebne do wybranej akcji."""
        selected_types = tuple(
            env_type for env_type in ("C5", "C6") if env_type in envelope_types
        )
        templates = {
            "C5": self.c5_tmpl_edit.text().strip(),
            "C6": self.c6_tmpl_edit.text().strip(),
        }
        missing_templates = [
            env_type
            for env_type in selected_types
            if not templates[env_type] or not Path(templates[env_type]).is_file()
        ]
        if missing_templates:
            QMessageBox.warning(
                self,
                "Brak szablonu",
                "Wskaż poprawny szablon: " + ", ".join(missing_templates) + ".",
            )
            return None

        self._save_sender_preferences()
        sender = (
            self.sender_name_edit.text().strip(),
            self.sender_street_edit.text().strip(),
            self.sender_city_edit.text().strip(),
        )
        if not all(sender):
            QMessageBox.warning(self, "Błąd", "Uzupełnij dane nadawcy.")
            return None
        return templates, sender

    def _generate_c5_envelopes(self):
        self._generate_envelopes(("C5",))

    def _generate_c6_envelopes(self):
        self._generate_envelopes(("C6",))

    def _generate_batch_envelopes(self):
        self._generate_envelopes(("C5", "C6"))

    def _generate_envelopes(self, envelope_types):
        """Generuje tylko wskazane rodzaje kopert i ich odpowiadające znaczki."""
        selected_types = tuple(
            env_type for env_type in ("C5", "C6") if env_type in envelope_types
        )
        if not selected_types:
            return

        targets = self._selected_targets(validate=True)
        if not targets:
            QMessageBox.warning(self, "Brak", "Zaznacz poprawnych adresatów.")
            return
        inputs = self._generation_inputs(selected_types)
        if not inputs:
            return
        templates, sender = inputs

        available = {
            "C5": sum(not stamp.get("used") for stamp in self.stamps_c5),
            "C6": sum(not stamp.get("used") for stamp in self.stamps_c6),
        }
        possible = min(available[env_type] for env_type in selected_types)
        requested_label = " + ".join(selected_types)
        if possible <= 0:
            missing = [
                env_type for env_type in selected_types if available[env_type] <= 0
            ]
            QMessageBox.critical(
                self,
                "Brak znaczków",
                "Brak wolnych znaczków: " + ", ".join(missing) + ".",
            )
            return
        if possible < len(targets):
            availability = ", ".join(
                f"{env_type}: {available[env_type]}"
                for env_type in selected_types
            )
            answer = QMessageBox.question(
                self,
                "Brakuje znaczków",
                f"Wolne znaczki ({availability}) wystarczą dla {possible} z "
                f"{len(targets)} adresatów ({requested_label}). Kontynuować?",
            )
            if answer != QMessageBox.StandardButton.Yes:
                return
            targets = targets[:possible]
        self._perform_generation(
            targets, templates, sender, selected_types, reuse=False
        )

    def _regenerate_envelopes(self):
        selected_types = ("C5", "C6")
        targets = self._selected_targets(validate=False)
        if not targets:
            QMessageBox.warning(self, "Brak", "Zaznacz adresatów.")
            return
        inputs = self._generation_inputs(selected_types)
        if not inputs:
            return
        templates, sender = inputs
        self._perform_generation(
            targets, templates, sender, selected_types, reuse=True
        )

    def _perform_generation(
        self, targets, templates, sender, envelope_types=("C5", "C6"), reuse=False
    ):
        selected_types = tuple(
            env_type for env_type in ("C5", "C6") if env_type in envelope_types
        )
        if not selected_types:
            return

        output = self._get_output_dir()
        single = self.chk_single_files.isChecked()
        temp = output if single else output / f"temp_{uuid.uuid4().hex[:8]}"
        temp.mkdir(parents=True, exist_ok=True)
        files_by_type = {"C5": [], "C6": []}
        shipments = self._load_shipments()
        suffix = "_KOREKTA" if reuse else ""
        from utils.docx_utils import generate_envelope_c5, generate_envelope_c6

        for order, (owner_index, address) in enumerate(targets):
            owner = self.owners[owner_index]
            name = self._get_formatted_name(owner)
            street, city = self._split_address(address)
            short = self._safe_filename(self._get_short_name(owner) + suffix)
            c5_stamp = c6_stamp = None

            if "C5" in selected_types:
                if reuse:
                    c5_record = next(
                        (
                            record
                            for record in shipments
                            if record.get("addressee") == name
                            and record.get("envelope_type") == "C5"
                        ),
                        None,
                    )
                    c5_stamp = next(
                        (
                            stamp
                            for stamp in self.stamps_c5
                            if c5_record
                            and stamp.get("barcode")
                            == c5_record.get("stamp_barcode")
                        ),
                        None,
                    )
                else:
                    c5_stamp = self._get_unused_stamp("C5")

            if "C6" in selected_types:
                if reuse:
                    c6_record = next(
                        (
                            record
                            for record in shipments
                            if name in record.get("addressee", "")
                            and record.get("envelope_type") == "C6"
                        ),
                        None,
                    )
                    c6_stamp = next(
                        (
                            stamp
                            for stamp in self.stamps_c6
                            if c6_record
                            and stamp.get("barcode")
                            == c6_record.get("stamp_barcode")
                        ),
                        None,
                    )
                else:
                    c6_stamp = self._get_unused_stamp("C6")

            required_stamps = {
                "C5": c5_stamp,
                "C6": c6_stamp,
            }
            if any(not required_stamps[env_type] for env_type in selected_types):
                break

            c5_ok = c6_ok = False
            c5_path = c6_path = None
            if "C5" in selected_types:
                c5_path = temp / f"C5_{order}_{short}.docx"
                c5_ok = generate_envelope_c5(
                    templates["C5"], str(c5_path), *sender,
                    name, street, city, c5_stamp.get("thumbnail_bytes"),
                )
                if c5_ok:
                    files_by_type["C5"].append(str(c5_path))

            if "C6" in selected_types:
                c6_path = temp / f"C6_{order}_{short}.docx"
                c6_ok = generate_envelope_c6(
                    templates["C6"], str(c6_path), *sender,
                    stamp_image_bytes=c6_stamp.get("thumbnail_bytes"),
                )
                if c6_ok:
                    files_by_type["C6"].append(str(c6_path))

            if not reuse:
                now = datetime.now().strftime("%Y-%m-%d %H:%M")
                if c5_ok:
                    c5_stamp.update(used=True, used_by=name)
                    owner["env_c5_generated"] = True
                    self.shipment_generated.emit({
                        "date": now, "addressee": name,
                        "addressee_street": street, "addressee_city": city,
                        "envelope_type": "C5",
                        "stamp_barcode": c5_stamp.get("barcode", ""),
                        "path": str(c5_path),
                    })
                if c6_ok:
                    c6_stamp.update(used=True, used_by=f"Zwrot od {name}")
                    owner["env_c6_generated"] = True
                    self.shipment_generated.emit({
                        "date": now, "addressee": f"Zwrot od {name}",
                        "addressee_street": sender[1],
                        "addressee_city": sender[2],
                        "envelope_type": "C6",
                        "stamp_barcode": c6_stamp.get("barcode", ""),
                        "path": str(c6_path),
                    })

        if not single:
            from utils.docx_utils import merge_docx_files

            stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            for env_type in selected_types:
                files = files_by_type[env_type]
                destination = output / (
                    f"Zbiorcze_Koperty_{env_type}{suffix}_{stamp}.docx"
                )
                if len(files) == 1:
                    shutil.copy(files[0], destination)
                elif len(files) > 1:
                    merge_docx_files(files, str(destination))
            shutil.rmtree(temp, ignore_errors=True)

        self._save_owners_state()
        self._save_stamp_status()
        self._refresh_stamps_list("C5")
        self._refresh_stamps_list("C6")
        self._update_stamps_info()
        self._refresh_owners_table()
        summary = "\n".join(
            f"Wygenerowano {env_type}: {len(files_by_type[env_type])}"
            for env_type in selected_types
        )
        QMessageBox.information(self, "Koniec", summary)

    # -------------------------------------------------------------- public
    def set_owners(self, owners_from_registry: list):
        self._add_unique_owners(owners_from_registry)

    def set_project(self, project: dict):
        self.active_project_path = project.get("path")
        self._load_owners_state()

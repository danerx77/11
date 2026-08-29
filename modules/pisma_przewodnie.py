"""
cover_letter.py – Zakładka generowania pism przewodnich
"""

import json
import re
import sys
import logging
from pathlib import Path
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QLineEdit,
    QGroupBox, QTextEdit, QDialog, QDialogButtonBox, QTableWidget,
    QTableWidgetItem, QFormLayout, QMessageBox, QFileDialog, QScrollArea, 
    QAbstractItemView, QCheckBox, QHeaderView, QComboBox, QListWidget, QListWidgetItem, QInputDialog
)
from PySide6.QtCore import Qt, Signal
from PySide6.QtGui import QColor, QFont

# Konfiguracja logera
def get_log_dir():
    if getattr(sys, 'frozen', False):
        return Path(sys.executable).parent.resolve() / 'dane'
    return Path(__file__).parent.parent.resolve() / 'dane'

log_dir = get_log_dir()
log_dir.mkdir(parents=True, exist_ok=True)
logging.basicConfig(
    filename=str(log_dir / 'error_log.txt'),
    level=logging.ERROR,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    encoding='utf-8'
)
logger = logging.getLogger(__name__)

def read_docx_text(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        lines = []
        for para in doc.paragraphs: lines.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs: lines.append(para.text)
        return '\n'.join(l for l in lines if l.strip())
    except Exception as e:
        logger.exception("Błąd podglądu pliku DOCX")
        return f'[Błąd odczytu: {e}]'

class NoWheelComboBox(QComboBox):
    def wheelEvent(self, e): e.ignore()

class ManageExamplesDialog(QDialog):
    def __init__(self, examples: list, title: str, parent=None):
        super().__init__(parent)
        self.setWindowTitle(f'Przykłady – {title}')
        self.setMinimumSize(500, 400)
        layout = QVBoxLayout(self)
        self.examples = list(examples)
        self.list_widget = QListWidget()
        for ex in self.examples: self.list_widget.addItem(ex)
        layout.addWidget(self.list_widget)
        add_row = QHBoxLayout()
        self.new_edit = QLineEdit()
        self.new_edit.setPlaceholderText('Nowy przykład...')
        add_row.addWidget(self.new_edit)
        btn_add = QPushButton('Dodaj')
        btn_add.clicked.connect(self._add)
        add_row.addWidget(btn_add)
        layout.addLayout(add_row)
        btn_row = QHBoxLayout()
        btn_edit = QPushButton('Edytuj zaznaczony')
        btn_edit.clicked.connect(self._edit)
        btn_row.addWidget(btn_edit)
        btn_del = QPushButton('Usuń zaznaczony')
        btn_del.clicked.connect(self._delete)
        btn_row.addWidget(btn_del)
        btn_row.addStretch()
        layout.addLayout(btn_row)
        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

    def _add(self):
        text = self.new_edit.text().strip()
        if text:
            self.examples.append(text)
            self.list_widget.addItem(text)
            self.new_edit.clear()

    def _delete(self):
        row = self.list_widget.currentRow()
        if row >= 0:
            self.examples.pop(row)
            self.list_widget.takeItem(row)

    def _edit(self):
        item = self.list_widget.currentItem()
        if item:
            old_text = item.text()
            new_text, ok = QInputDialog.getText(self, "Edycja", "Zmień tekst przykładu:", QLineEdit.EchoMode.Normal, old_text)
            if ok and new_text.strip():
                row = self.list_widget.row(item)
                self.examples[row] = new_text.strip()
                item.setText(new_text.strip())
                
    def get_examples(self): return self.examples

class OwnershipPhraseDialog(QDialog):
    def __init__(self, config: dict, parent=None):
        super().__init__(parent)
        self.config = config
        self.setWindowTitle("Edytor Szablonów Formuł Własności")
        self.setMinimumWidth(600)
        self.setMinimumHeight(650)
        
        self.default_phrases = {
            "couple_sole_sing": "której są Państwo właścicielami,",
            "couple_sole_plur": "których są Państwo właścicielami,",
            "couple_joint_sing": "której są Państwo współwłaścicielami,",
            "couple_joint_plur": "których są Państwo współwłaścicielami,",
            
            "female_sole_sing": "której jest Pani właścicielką,",
            "female_sole_plur": "których jest Pani właścicielką,",
            "female_joint_sing": "której jest Pani współwłaścicielką,",
            "female_joint_plur": "których jest Pani współwłaścicielką,",
            
            "male_sole_sing": "której jest Pan właścicielem,",
            "male_sole_plur": "których jest Pan właścicielem,",
            "male_joint_sing": "której jest Pan współwłaścicielem,",
            "male_joint_plur": "których jest Pan współwłaścicielem,",
            
            "company_sole_sing": "której są Państwo właścicielem,",
            "company_sole_plur": "których są Państwo właścicielem,",
            "company_joint_sing": "której są Państwo współwłaścicielem,",
            "company_joint_plur": "których są Państwo współwłaścicielem,"
        }
        
        self.phrases = self.config.get('ownership_phrases', self.default_phrases)
        # Zapewnienie, że nowe klucze "company" istnieją w załadowanej konfiguracji
        for k, v in self.default_phrases.items():
            if k not in self.phrases:
                self.phrases[k] = v
                
        self.edits = {}
        
        layout = QVBoxLayout(self)
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        inner = QWidget()
        form = QFormLayout(inner)
        
        lbl_info = QLabel("Edytuj zwroty używane automatycznie w Piśmie Przewodnim:")
        lbl_info.setStyleSheet("font-weight: bold; margin-bottom: 10px;")
        form.addRow(lbl_info)
        
        labels_map = {
            "couple_sole_sing": "PARA | Właściciel | 1 działka:",
            "couple_sole_plur": "PARA | Właściciel | Wiele działek:",
            "couple_joint_sing": "PARA | Współwłaściciel | 1 działka:",
            "couple_joint_plur": "PARA | Współwłaściciel | Wiele działek:",
            
            "female_sole_sing": "KOBIETA | Właściciel | 1 działka:",
            "female_sole_plur": "KOBIETA | Właściciel | Wiele działek:",
            "female_joint_sing": "KOBIETA | Współwłaściciel | 1 działka:",
            "female_joint_plur": "KOBIETA | Współwłaściciel | Wiele działek:",
            
            "male_sole_sing": "MĘŻCZYZNA | Właściciel | 1 działka:",
            "male_sole_plur": "MĘŻCZYZNA | Właściciel | Wiele działek:",
            "male_joint_sing": "MĘŻCZYZNA | Współwłaściciel | 1 działka:",
            "male_joint_plur": "MĘŻCZYZNA | Współwłaściciel | Wiele działek:",
            
            "company_sole_sing": "FIRMA/INSTYTUCJA | Właściciel | 1 działka:",
            "company_sole_plur": "FIRMA/INSTYTUCJA | Właściciel | Wiele działek:",
            "company_joint_sing": "FIRMA/INSTYTUCJA | Współwłaściciel | 1 działka:",
            "company_joint_plur": "FIRMA/INSTYTUCJA | Współwłaściciel | Wiele działek:"
        }
        
        for key, label in labels_map.items():
            edit = QLineEdit(self.phrases.get(key, self.default_phrases[key]))
            self.edits[key] = edit
            form.addRow(label, edit)
            
        scroll.setWidget(inner)
        layout.addWidget(scroll)
        
        btns = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        btns.accepted.connect(self.accept)
        btns.rejected.connect(self.reject)
        layout.addWidget(btns)
        
    def accept(self):
        new_phrases = {}
        for key, edit in self.edits.items():
            new_phrases[key] = edit.text().strip()
        self.config['ownership_phrases'] = new_phrases
        super().accept()

class CoverLetterWidget(QWidget):
    owners_changed = Signal(list)

    def __init__(self, config: dict, examples: dict, save_callback=None, parent=None):
        super().__init__(parent)
        self.config = config
        self.examples = examples
        self.save_callback = save_callback
        self.owners = []
        self.parcels = []
        self.checked_keys = set()
        self.parcel_groups = {}
        self.active_project_path = ''
        self.setAcceptDrops(True)
        self._build_ui()

    def showEvent(self, event):
        super().showEvent(event)
        cl = self.config.get('cover_letter_template', '')
        if cl: self.template_edit.setText(cl)

    def dragEnterEvent(self, e):
        if e.mimeData().hasUrls(): e.acceptProposedAction()

    def dropEvent(self, e):
        for url in e.mimeData().urls():
            if url.isLocalFile():
                path = url.toLocalFile()
                lower = path.lower()
                if lower.endswith('.txt'):
                    self._import_group_txt(path)
                elif lower.endswith('.docx'):
                    self.template_edit.setText(path)
                break

    def _build_ui(self):
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        inner = QWidget()
        main_v = QVBoxLayout(inner)
        main_v.setSpacing(10)

        hdr = QLabel('📬 Generator Pism Przewodnich')
        hdr.setStyleSheet('font-size:16px; font-weight:700;')
        main_v.addWidget(hdr)

        meta_box = QGroupBox('Nagłówek pisma')
        meta_form = QFormLayout(meta_box)
        self.date_edit = QLineEdit()
        self.date_edit.setPlaceholderText('np. 20.10.2026')
        meta_form.addRow('Data sporządzenia:', self.date_edit)
        self.sender_place_edit = QLineEdit(self.config.get('sender', {}).get('city', ''))
        self.sender_place_edit.setPlaceholderText('np. Gdańsk')
        meta_form.addRow('Miejscowość druku:', self.sender_place_edit)
        main_v.addWidget(meta_box)

        owner_box = QGroupBox('Lista Właścicieli (Wybierz aby edytować dane do pisma)')
        owner_layout = QVBoxLayout(owner_box)
        search_row = QHBoxLayout()
        search_row.addWidget(QLabel('Wyszukaj (np. Jan Kowalski, 123/1):'))
        self.search_owners_edit = QLineEdit()
        self.search_owners_edit.textChanged.connect(self._refresh_owners_table)
        search_row.addWidget(self.search_owners_edit)
        owner_layout.addLayout(search_row)

        chk_layout = QHBoxLayout()
        self.chk_hide_generated = QCheckBox("Ukryj osoby, dla których wygenerowano pismo")
        self.chk_hide_generated.setChecked(False)
        self.chk_hide_generated.stateChanged.connect(self._refresh_owners_table)
        chk_layout.addWidget(self.chk_hide_generated)
        self.chk_show_only_generated = QCheckBox('Pokaż tylko wygenerowane')
        self.chk_show_only_generated.stateChanged.connect(self._refresh_owners_table)
        chk_layout.addWidget(self.chk_show_only_generated)
        btn_sel_all = QPushButton('☑️ Zaznacz wszystkie ptaszki')
        btn_sel_all.clicked.connect(self._check_all_visible)
        chk_layout.addWidget(btn_sel_all)
        btn_unsel_all = QPushButton('⬜ Odznacz wszystkie ptaszki')
        btn_unsel_all.clicked.connect(self._uncheck_all_visible)
        chk_layout.addWidget(btn_unsel_all)
        chk_layout.addStretch()

        owner_layout.addLayout(chk_layout)

        group_layout = QHBoxLayout()
        group_layout.addWidget(QLabel('Grupa działek:'))
        self.group_combo = QComboBox()
        self.group_combo.setMinimumWidth(180)
        self.group_combo.currentTextChanged.connect(self._on_group_changed)
        group_layout.addWidget(self.group_combo)
        self.chk_show_only_group = QCheckBox('Pokaż tylko grupę')
        self.chk_show_only_group.stateChanged.connect(self._refresh_owners_table)
        group_layout.addWidget(self.chk_show_only_group)
        self.chk_exclude_grouped_from_all = QCheckBox('Wszystkie bez działek z grup')
        self.chk_exclude_grouped_from_all.setToolTip('Gdy wybierzesz "Wszystkie działki", działki użyte w innych grupach nie będą pokazywane.')
        self.chk_exclude_grouped_from_all.stateChanged.connect(self._refresh_owners_table)
        group_layout.addWidget(self.chk_exclude_grouped_from_all)
        btn_group_create = QPushButton('➕ Utwórz z ptaszków')
        btn_group_create.clicked.connect(self._create_group_from_checked)
        group_layout.addWidget(btn_group_create)
        btn_group_apply = QPushButton('☑️ Zaznacz grupę')
        btn_group_apply.clicked.connect(self._apply_selected_group)
        group_layout.addWidget(btn_group_apply)
        btn_group_import = QPushButton('📥 Import TXT jako grupa')
        btn_group_import.clicked.connect(self._import_group_txt)
        group_layout.addWidget(btn_group_import)
        btn_group_delete = QPushButton('🗑 Usuń grupę')
        btn_group_delete.clicked.connect(self._delete_selected_group)
        group_layout.addWidget(btn_group_delete)
        group_layout.addStretch()
        owner_layout.addLayout(group_layout)

        self.table_owners = QTableWidget(0, 5)
        self.table_owners.setHorizontalHeaderLabels(['✓', 'Wygenerowano', 'Właściciel / Opis', 'Działki', 'Adres'])
        self.table_owners.setMinimumHeight(280)
        
        self.table_owners.horizontalHeader().setSectionsMovable(True)
        table_state_cover_hex = self.config.get('table_state_cover', '')
        if table_state_cover_hex:
            from PySide6.QtCore import QByteArray
            self.table_owners.horizontalHeader().restoreState(QByteArray.fromHex(table_state_cover_hex.encode()))
        header = self.table_owners.horizontalHeader()
        for col in range(self.table_owners.columnCount()):
            header.setSectionResizeMode(col, QHeaderView.ResizeMode.Interactive)
        self.table_owners.setColumnWidth(0, 60)
        self.table_owners.setColumnWidth(1, 115)
        self.table_owners.setColumnWidth(2, 250)
        self.table_owners.setColumnWidth(3, 150)
        self.table_owners.setColumnWidth(4, 240)
        self.table_owners.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        
        if not table_state_cover_hex:
            self.table_owners.setColumnWidth(2, 250)
            self.table_owners.setColumnWidth(3, 150)
            
        self.table_owners.horizontalHeader().sectionResized.connect(lambda *args: self.config.update({'table_state_cover': self.table_owners.horizontalHeader().saveState().toHex().data().decode()}))
        self.table_owners.horizontalHeader().sectionMoved.connect(lambda *args: self.config.update({'table_state_cover': self.table_owners.horizontalHeader().saveState().toHex().data().decode()}))
        self.table_owners.setAlternatingRowColors(True)
        self.table_owners.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self.table_owners.setSelectionMode(QAbstractItemView.SelectionMode.ExtendedSelection)
        self.table_owners.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        self.table_owners.itemSelectionChanged.connect(self._on_owner_selected)
        self.table_owners.itemChanged.connect(self._on_check_changed)
        
        owner_layout.addWidget(self.table_owners)

        manual_btns = QHBoxLayout()
        btn_mark_gen = QPushButton('✅ Oznacz wybrane wiersze jako Gotowe')
        btn_mark_gen.clicked.connect(self._mark_selected_as_generated)
        manual_btns.addWidget(btn_mark_gen)
        btn_mark_ungen = QPushButton('❌ Oznacz wybrane wiersze jako Brak')
        btn_mark_ungen.clicked.connect(self._mark_selected_as_ungenerated)
        manual_btns.addWidget(btn_mark_ungen)
        manual_btns.addStretch()
        owner_layout.addLayout(manual_btns)

        addr_form = QFormLayout()
        self.addr_name_edit = QLineEdit()
        addr_form.addRow('Imię i Nazwisko (Adresat):', self.addr_name_edit)
        self.addr_street_edit = QLineEdit()
        addr_form.addRow('Ulica:', self.addr_street_edit)
        self.addr_city_edit = QLineEdit()
        addr_form.addRow('Kod pocztowy i Miejscowość:', self.addr_city_edit)
        owner_layout.addLayout(addr_form)
        main_v.addWidget(owner_box)

        content_box = QGroupBox('Treść pisma')
        content_layout = QVBoxLayout(content_box)
        parcels_form = QFormLayout()
        self.location_edit = QLineEdit()
        parcels_form.addRow('Miejscowość działki:', self.location_edit)
        self.street_dz_edit = QLineEdit()
        parcels_form.addRow('Ulica działki:', self.street_dz_edit)
        
        self.subject_edit = QLineEdit()
        parcels_form.addRow('Temat:', self.subject_edit)
        h_task_bud = QHBoxLayout()
        self.task_budowa_combo = NoWheelComboBox()
        self.task_budowa_combo.currentTextChanged.connect(lambda t: self.task_budowa_edit.setText(t) if t else None)
        h_task_bud.addWidget(self.task_budowa_combo, 1)
        btn_manage_task_bud = QPushButton('✏ Przykłady')
        btn_manage_task_bud.clicked.connect(lambda: self._manage_task_examples('budowa'))
        h_task_bud.addWidget(btn_manage_task_bud)
        self.task_budowa_edit = QLineEdit()
        h_task_bud.addWidget(self.task_budowa_edit, 2)
        parcels_form.addRow('Zadanie budowa:', h_task_bud)

        h_task_dem = QHBoxLayout()
        self.task_demontaz_combo = NoWheelComboBox()
        self.task_demontaz_combo.currentTextChanged.connect(lambda t: self.task_demontaz_edit.setText(t) if t else None)
        h_task_dem.addWidget(self.task_demontaz_combo, 1)
        btn_manage_task_dem = QPushButton('✏ Przykłady')
        btn_manage_task_dem.clicked.connect(lambda: self._manage_task_examples('demontaz'))
        h_task_dem.addWidget(btn_manage_task_dem)
        self.task_demontaz_edit = QLineEdit()
        h_task_dem.addWidget(self.task_demontaz_edit, 2)
        parcels_form.addRow('Zadanie demontaż:', h_task_dem)
        
        self.parcels_edit = QLineEdit()
        self.parcels_edit.setPlaceholderText('np. 453/5, 453/6')
        parcels_form.addRow('Działki ogółem:', self.parcels_edit)
        self.parcels_budowa_edit = QLineEdit()
        parcels_form.addRow('Działki budowa:', self.parcels_budowa_edit)
        self.parcels_demontaz_edit = QLineEdit()
        parcels_form.addRow('Działki demontaż:', self.parcels_demontaz_edit)
        
        content_layout.addLayout(parcels_form)
        own_row = QHBoxLayout()
        own_row.addWidget(QLabel('Formuła własności (Edytowalna):'))
        self.ownership_edit = QLineEdit()
        self.ownership_edit.setPlaceholderText('Możesz tu dowolnie zmienić wygenerowaną formułę...')
        own_row.addWidget(self.ownership_edit, 2)
        btn_auto_own = QPushButton('↩ Generuj formę gramatyczną')
        btn_auto_own.clicked.connect(self._auto_ownership)
        own_row.addWidget(btn_auto_own)
        
        btn_edit_phrases = QPushButton('⚙️ Edytuj szablony formuł')
        btn_edit_phrases.clicked.connect(self._edit_ownership_phrases)
        own_row.addWidget(btn_edit_phrases)
        
        content_layout.addLayout(own_row)
        main_v.addWidget(content_box)

        tmpl_box = QGroupBox('Szablon')
        tmpl_layout = QHBoxLayout(tmpl_box)
        self.template_edit = QLineEdit()
        self.template_edit.setPlaceholderText('Ścieżka do szablonu Pismo przewodnie .docx')
        tmpl_layout.addWidget(self.template_edit)
        btn_browse = QPushButton('📂 Wybierz...')
        btn_browse.clicked.connect(self._browse_template)
        tmpl_layout.addWidget(btn_browse)
        btn_default = QPushButton('↩ Domyślny')
        btn_default.clicked.connect(self._set_default_template)
        tmpl_layout.addWidget(btn_default)
        main_v.addWidget(tmpl_box)

        gen_box = QGroupBox('Generuj')
        gen_v = QVBoxLayout(gen_box)
        gen_row = QHBoxLayout()
        self.btn_preview = QPushButton('👁 Podgląd tekstu')
        self.btn_preview.clicked.connect(self._preview)
        gen_row.addWidget(self.btn_preview)
        self.btn_generate = QPushButton('💾 Generuj dla formularza')
        self.btn_generate.clicked.connect(self._generate)
        gen_row.addWidget(self.btn_generate)
        self.btn_generate_sel = QPushButton('⚙️ Generuj dla ZAZNACZONYCH (z ptaszkiem)')
        self.btn_generate_sel.clicked.connect(self._generate_selected)
        gen_row.addWidget(self.btn_generate_sel)
        self.btn_generate_all = QPushButton('⚡ GENERUJ WSZYSTKIE BEZ PTASZKA')
        self.btn_generate_all.setObjectName('btn_accent')
        self.btn_generate_all.clicked.connect(self._generate_all)
        gen_row.addWidget(self.btn_generate_all)
        gen_v.addLayout(gen_row)
        self.preview_text = QTextEdit()
        self.preview_text.setReadOnly(True)
        self.preview_text.setMaximumHeight(140)
        gen_v.addWidget(self.preview_text)
        main_v.addWidget(gen_box)
        main_v.addStretch()

        scroll.setWidget(inner)
        outer = QVBoxLayout(self)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.addWidget(scroll)
        self._update_task_examples()

    def _update_task_examples(self):
        self.task_budowa_combo.blockSignals(True)
        self.task_budowa_combo.clear()
        self.task_budowa_combo.addItem('')
        for ex in self.examples.get('task_types_budowa', []): 
            self.task_budowa_combo.addItem(ex)
        self.task_budowa_combo.blockSignals(False)

        self.task_demontaz_combo.blockSignals(True)
        self.task_demontaz_combo.clear()
        self.task_demontaz_combo.addItem('')
        for ex in self.examples.get('task_types_demontaz', []): 
            self.task_demontaz_combo.addItem(ex)
        self.task_demontaz_combo.blockSignals(False)

    def _manage_task_examples(self, task_type: str):
        key = 'task_types_budowa' if task_type == 'budowa' else 'task_types_demontaz'
        title = 'Zadanie Budowa' if task_type == 'budowa' else 'Zadanie Demontaż'
        dlg = ManageExamplesDialog(self.examples.get(key, []), title, self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            self.examples[key] = dlg.get_examples()
            self._update_task_examples()
            if self.save_callback: self.save_callback()

    def _edit_ownership_phrases(self):
        dlg = OwnershipPhraseDialog(self.config, self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            self._auto_ownership()
            if self.save_callback: self.save_callback()

    def _check_all_visible(self):
        for r in range(self.table_owners.rowCount()):
            if not self.table_owners.isRowHidden(r):
                it = self.table_owners.item(r, 0)
                if it: it.setCheckState(Qt.CheckState.Checked)

    def _uncheck_all_visible(self):
        for r in range(self.table_owners.rowCount()):
            it = self.table_owners.item(r, 0)
            if it: it.setCheckState(Qt.CheckState.Unchecked)

    def _cover_done(self, owner, specific_addr):
        by_addr = owner.get('cover_generated_by_addr')
        if isinstance(by_addr, dict):
            return bool(by_addr.get(specific_addr or '', False))
        return bool(owner.get('cover_generated', False))

    def _set_cover_done(self, owner, specific_addr, value):
        by_addr = owner.get('cover_generated_by_addr')
        if not isinstance(by_addr, dict):
            by_addr = {}
            owner['cover_generated_by_addr'] = by_addr
        by_addr[specific_addr or ''] = bool(value)
        owner['cover_generated'] = any(bool(v) for v in by_addr.values())

    def _row_key(self, idx, specific_addr):
        return f"{idx}|{specific_addr or ''}"

    def _on_check_changed(self, item):
        if item.column() != 0:
            return
        key = item.data(Qt.ItemDataRole.UserRole)
        if not key:
            return
        if item.checkState() == Qt.CheckState.Checked:
            self.checked_keys.add(key)
        else:
            self.checked_keys.discard(key)

    def _update_generated_list(self):
        if not hasattr(self, 'generated_list'):
            return
        self.generated_list.clear()
        for o in self.owners:
            name = o.get('full_name') or o.get('name_plural') or ''
            parcels = ', '.join([str(p.get('number', p)) if isinstance(p, dict) else str(p) for p in o.get('parcels', [])])
            if o.get('cover_generated'):
                self.generated_list.addItem(f"{name} | {parcels} | Pismo przewodne")

    def _toggle_generated_list(self):
        self.generated_list.setVisible(not self.generated_list.isVisible())

    def _owner_parcels(self, owner):
        return {str(p.get('number', p)) if isinstance(p, dict) else str(p) for p in owner.get('parcels', [])}

    def _set_group_fields(self, group: dict | None):
        """Ustawia pola zadań dla wybranej grupy. Dla 'Wszystkie działki' czyści pola."""
        if not group:
            self.subject_edit.clear()
            self.task_budowa_edit.clear()
            self.task_demontaz_edit.clear()
            return
        self.subject_edit.setText(group.get('subject', '') or '')
        self.task_budowa_edit.setText(group.get('budowa', '') or '')
        self.task_demontaz_edit.setText(group.get('demontaz', '') or '')

    def _on_group_changed(self, *_args):
        name = self.group_combo.currentText() if hasattr(self, 'group_combo') else ''
        if name == 'Wszystkie działki' or not name:
            self._set_group_fields(None)
        else:
            self._set_group_fields(self.parcel_groups.get(name, {}))
        if hasattr(self, 'chk_show_only_group') and self.chk_show_only_group.isChecked():
            self._refresh_owners_table()

    def _groups_state_file(self):
        project_path = self.active_project_path or self.config.get('last_project_path', '')
        return Path(project_path) / 'cover_groups.json' if project_path else None

    def _save_groups(self):
        path = self._groups_state_file()
        if not path: return
        try:
            path.parent.mkdir(parents=True, exist_ok=True)
            with open(path, 'w', encoding='utf-8') as f:
                json.dump(self.parcel_groups, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def _load_groups(self):
        path = self._groups_state_file()
        self.parcel_groups = {}
        if path and path.exists():
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                if isinstance(data, dict):
                    self.parcel_groups = data
            except Exception:
                self.parcel_groups = {}
        self._refresh_group_combo()

    def _refresh_group_combo(self):
        if not hasattr(self, 'group_combo'):
            return
        current = self.group_combo.currentText()
        self.group_combo.blockSignals(True)
        self.group_combo.clear()
        self.group_combo.addItem('Wszystkie działki')
        self.group_combo.addItems(sorted(self.parcel_groups.keys()))
        if current:
            idx = self.group_combo.findText(current)
            if idx >= 0: self.group_combo.setCurrentIndex(idx)
        self.group_combo.blockSignals(False)
        self._on_group_changed()

    def _owner_group_key(self, owner, specific_addr):
        parcels = ','.join(sorted(self._owner_parcels(owner)))
        name = owner.get('full_name') or owner.get('name_plural') or owner.get('name_separate') or owner.get('last_name', '')
        return f"{name}|{specific_addr or ''}|{parcels}"

    def _selected_group_owner_keys(self):
        if not hasattr(self, 'group_combo'):
            return set()
        name = self.group_combo.currentText()
        group = self.parcel_groups.get(name, {})
        return set(group.get('owner_keys', []))

    def _selected_group_parcels(self):
        if not hasattr(self, 'group_combo'):
            return set()
        name = self.group_combo.currentText()
        if name == 'Wszystkie działki':
            all_nums = set(self._all_project_parcels())
            if hasattr(self, 'chk_exclude_grouped_from_all') and self.chk_exclude_grouped_from_all.isChecked():
                used = set()
                for group in self.parcel_groups.values():
                    used.update(group.get('parcels', []))
                all_nums -= used
            return all_nums
        group = self.parcel_groups.get(name, {})
        return set(group.get('parcels', []))

    def _all_project_parcels(self):
        nums = set()
        for p in self.parcels:
            if isinstance(p, dict) and p.get('number'):
                nums.add(str(p.get('number')))
        for o in self.owners:
            nums.update(self._owner_parcels(o))
        return nums

    def _ask_group_descriptions(self):
        dlg = QDialog(self)
        dlg.setWindowTitle('Temat i zadania dla grupy')
        layout = QFormLayout(dlg)

        subject_edit = QLineEdit(self.subject_edit.text().strip() if hasattr(self, 'subject_edit') else '')
        layout.addRow('Temat – wpisz dla tej grupy:', subject_edit)

        bud_combo = NoWheelComboBox()
        bud_combo.setEditable(True)
        bud_combo.addItem('')
        for ex in self.examples.get('task_types_budowa', []):
            bud_combo.addItem(ex)
        layout.addRow('Zadanie Budowa – wybierz z listy lub wpisz:', bud_combo)

        dem_combo = NoWheelComboBox()
        dem_combo.setEditable(True)
        dem_combo.addItem('')
        for ex in self.examples.get('task_types_demontaz', []):
            dem_combo.addItem(ex)
        layout.addRow('Zadanie Demontaż – wybierz z listy lub wpisz:', dem_combo)

        info = QLabel('Możesz wybrać istniejące zadanie z listy albo wpisać własne dla tej grupy.')
        info.setStyleSheet('color: gray; font-size: 11px;')
        layout.addRow('', info)

        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(dlg.accept)
        buttons.rejected.connect(dlg.reject)
        layout.addRow(buttons)

        if dlg.exec() != QDialog.DialogCode.Accepted:
            return None
        return subject_edit.text().strip(), bud_combo.currentText().strip(), dem_combo.currentText().strip()


    def _create_group_from_checked(self):
        parcels = set()
        owner_keys = set()
        for key in self.checked_keys:
            try:
                raw = str(key).split('|', 1)
                idx = int(raw[0])
                specific_addr = raw[1] if len(raw) > 1 else ''
                if 0 <= idx < len(self.owners):
                    owner = self.owners[idx]
                    parcels.update(self._owner_parcels(owner))
                    owner_keys.add(self._owner_group_key(owner, specific_addr))
            except Exception:
                pass
        if not owner_keys:
            return QMessageBox.warning(self, 'Brak', 'Najpierw zaznacz ptaszkiem konkretne osoby/adresy do grupy.')
        name, ok = QInputDialog.getText(self, 'Nowa grupa', 'Nazwa grupy:')
        if not ok or not name.strip(): return
        descs = self._ask_group_descriptions()
        if descs is None: return
        self.parcel_groups[name.strip()] = {'parcels': sorted(parcels), 'owner_keys': sorted(owner_keys), 'subject': descs[0], 'budowa': descs[1], 'demontaz': descs[2]}
        self._save_groups()
        self._refresh_group_combo()
        self.group_combo.setCurrentText(name.strip())


    def _apply_selected_group(self):
        name = self.group_combo.currentText() if hasattr(self, 'group_combo') else ''
        if name == 'Wszystkie działki':
            group = {'parcels': sorted(self._all_project_parcels()), 'budowa': '', 'demontaz': ''}
        else:
            group = self.parcel_groups.get(name)
        if not group: return

        owner_keys = set(group.get('owner_keys', []))
        wanted = set(group.get('parcels', []))
        for idx, o in enumerate(self.owners):
            addresses = [o.get('address', '')] + ([o.get('address_2')] if o.get('address_2') else [])
            for addr in addresses:
                if owner_keys:
                    if self._owner_group_key(o, addr) in owner_keys:
                        self.checked_keys.add(self._row_key(idx, addr))
                elif self._owner_parcels(o) & wanted:
                    self.checked_keys.add(self._row_key(idx, addr))
        self._set_group_fields(None if name == 'Wszystkie działki' else group)
        self._refresh_owners_table()


    def _delete_selected_group(self):
        name = self.group_combo.currentText() if hasattr(self, 'group_combo') else ''
        if name and name != 'Wszystkie działki' and name in self.parcel_groups:
            del self.parcel_groups[name]
            self._save_groups()
            self._refresh_group_combo()

    def _import_group_txt(self, path: str = ''):
        if not path:
            path, _ = QFileDialog.getOpenFileName(self, 'Import grupy działek TXT', '', 'Pliki TXT (*.txt);;Wszystkie pliki (*.*)')
        if not path: return
        try:
            text = Path(path).read_text(encoding='utf-8')
        except UnicodeDecodeError:
            text = Path(path).read_text(encoding='cp1250')
        from utils.pdf_utils import parse_parcel_list_text
        parsed = parse_parcel_list_text(text)
        lines = []
        for cat in ['demolition', 'construction', 'connection', 'full']:
            lines.extend(parsed.get(cat, []))
        # deduplikacja z zachowaniem kolejności
        seen = set()
        lines = [x for x in lines if not (x in seen or seen.add(x))]
        if not lines:
            QMessageBox.warning(self, 'Brak działek', 'Nie znaleziono numerów działek w pliku TXT.')
            return
        group_name, ok = QInputDialog.getText(self, 'Nazwa grupy', 'Podaj nazwę grupy:', QLineEdit.EchoMode.Normal, Path(path).stem)
        if not ok or not group_name.strip(): return
        descs = self._ask_group_descriptions()
        if descs is None: return
        self.parcel_groups[group_name.strip()] = {'parcels': lines, 'subject': descs[0], 'budowa': descs[1], 'demontaz': descs[2]}
        self._save_groups()
        self._refresh_group_combo()
        self.group_combo.setCurrentText(group_name.strip())
        self._on_group_changed()
        self._apply_selected_group()


    def set_owners(self, owners: list):
        self.owners = owners
        self._refresh_group_combo()
        self._refresh_owners_table()

    def _refresh_owners_table(self):
        self.table_owners.blockSignals(True)
        self.table_owners.setRowCount(0)
        hide_generated = self.chk_hide_generated.isChecked()
        filter_text = self.search_owners_edit.text().strip().lower()
        
        for idx, o in enumerate(self.owners):
            addresses = [o.get('address', '')]
            if o.get('address_2'): addresses.append(o.get('address_2'))
                
            for specific_addr in addresses:
                is_generated = self._cover_done(o, specific_addr)
                if hide_generated and is_generated: continue
                if hasattr(self, 'chk_show_only_generated') and self.chk_show_only_generated.isChecked() and not is_generated: continue
                
                fmt = self.config.get('couple_format_cover', 0)
                display_name = o.get('name_plural', o.get('full_name', '')) if fmt == 0 else o.get('name_separate', o.get('full_name', ''))
                parcels_str = ', '.join([str(p.get('number', p)) if isinstance(p, dict) else str(p) for p in o.get('parcels', [])])
                if hasattr(self, 'chk_show_only_group') and self.chk_show_only_group.isChecked():
                    wanted_owner_keys = self._selected_group_owner_keys()
                    if wanted_owner_keys:
                        if self._owner_group_key(o, specific_addr) not in wanted_owner_keys:
                            continue
                    else:
                        wanted_group = self._selected_group_parcels()
                        owner_nums = {str(p.get('number', p)) if isinstance(p, dict) else str(p) for p in o.get('parcels', [])}
                        if wanted_group and not (owner_nums & wanted_group):
                            continue
                if filter_text:
                    search_target = f"{display_name} {specific_addr} {parcels_str}".lower()
                    if filter_text not in search_target: continue
                flags = []
                if o.get('is_dead'): flags.append("[ZMARŁY/A]")
                elif o.get('is_institution'): flags.append("[INSTYTUCJA]")
                elif o.get('is_church'): flags.append("[PARAFIA]")
                elif o.get('is_spolka'): flags.append("[SPÓŁKA]")
                elif o.get('is_company'): flags.append("[FIRMA]")
                else:
                    if not specific_addr.strip(): flags.append("[BRAK ADRESU]")
                    elif not re.search(r'\d{2}-\d{3}', specific_addr): flags.append("[BŁĄD KODU]")
                    
                flag_str = " ".join(flags) + " " if flags else ""
                
                row = self.table_owners.rowCount()
                self.table_owners.insertRow(row)
                
                it_chk = QTableWidgetItem()
                it_chk.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled | Qt.ItemFlag.ItemIsSelectable)
                key = self._row_key(idx, specific_addr)
                it_chk.setData(Qt.ItemDataRole.UserRole, key)
                it_chk.setCheckState(Qt.CheckState.Checked if key in self.checked_keys else Qt.CheckState.Unchecked)
                self.table_owners.setItem(row, 0, it_chk)
                
                txt_gen = "TAK" if is_generated else "NIE"
                color_gen = QColor("#2ecc71") if is_generated else QColor("#e74c3c")
                it_gen = QTableWidgetItem(txt_gen)
                it_gen.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                it_gen.setForeground(color_gen)
                it_gen.setFont(QFont('', -1, QFont.Weight.Bold))
                self.table_owners.setItem(row, 1, it_gen)
                
                it_name = QTableWidgetItem(f"{flag_str}{display_name}")
                it_name.setData(Qt.ItemDataRole.UserRole, (idx, specific_addr))
                if o.get('is_dead'): it_name.setForeground(QColor("#e74c3c"))
                elif o.get('is_institution'): it_name.setForeground(QColor("#9b5de5"))
                elif o.get('is_spolka'): it_name.setForeground(QColor("#2980b9"))
                elif o.get('is_company'): it_name.setForeground(QColor("#3498db"))
                elif o.get('is_church'): it_name.setForeground(QColor("#f39c12"))
                elif flags: it_name.setForeground(QColor("#e67e22"))
                else: it_name.setForeground(QColor("#2ecc71"))
                self.table_owners.setItem(row, 2, it_name)
                
                it_parcels = QTableWidgetItem(parcels_str)
                self.table_owners.setItem(row, 3, it_parcels)
                it_addr = QTableWidgetItem(specific_addr)
                self.table_owners.setItem(row, 4, it_addr)
                
        self.table_owners.blockSignals(False)

    def _mark_selected_as_generated(self):
        sel = self.table_owners.selectionModel().selectedRows()
        if not sel: return QMessageBox.warning(self, "Brak", "Wybierz wiersze z tabeli (zaznacz na niebiesko).")
        for i in sel:
            data = self.table_owners.item(i.row(), 2).data(Qt.ItemDataRole.UserRole)
            if data:
                idx, _ = data
                self._set_cover_done(self.owners[idx], _, True)
        self.owners_changed.emit(self.owners)
        self._refresh_owners_table()

    def _mark_selected_as_ungenerated(self):
        sel = self.table_owners.selectionModel().selectedRows()
        if not sel: return QMessageBox.warning(self, "Brak", "Wybierz wiersze z tabeli (zaznacz na niebiesko).")
        for i in sel:
            data = self.table_owners.item(i.row(), 2).data(Qt.ItemDataRole.UserRole)
            if data:
                idx, _ = data
                self._set_cover_done(self.owners[idx], _, False)
        self.owners_changed.emit(self.owners)
        self._refresh_owners_table()

    def set_parcels(self, parcels: list): self.parcels = parcels
    def set_project(self, project: dict):
        self.active_project_path = project.get('path', '')
        self._load_groups()

    def _is_valid_for_gen(self, o: dict, specific_addr: str) -> tuple[bool, str]:
        if o.get('is_dead'): return False, "Osoba zmarła"
        if o.get('is_institution') or o.get('is_church'): return False, "Instytucja/Gmina/Parafia"
        if not specific_addr.strip(): return False, "Brak adresu"
        if not re.search(r'\d{2}-\d{3}', specific_addr): return False, "Brak kodu pocztowego"
        return True, "OK"

    def _get_short_name(self, fn: str, ln: str) -> str:
        if ' i ' in fn.lower():
            names = [n[0].upper() for n in fn.replace(' i ', ' ').split() if n.lower() != 'i']
            return f"{'.'.join(names)}.{ln}"
        return f"{fn[0].upper() if fn else ''}.{ln}"

    def _on_owner_selected(self):
        sel = self.table_owners.selectedItems()
        if not sel: return
        
        row = sel[0].row()
        data = self.table_owners.item(row, 2).data(Qt.ItemDataRole.UserRole)
        if not data: return
        idx, specific_addr = data
        o = self.owners[idx]
        
        fmt = self.config.get('couple_format_cover', 0)
        display_name = o.get('name_plural', o.get('full_name', '')) if fmt == 0 else o.get('name_separate', o.get('full_name', ''))
        
        self.addr_name_edit.setText(display_name)
        
        addr_parts = specific_addr.split(',')
        self.addr_street_edit.setText(addr_parts[0].strip() if addr_parts else '')
        if len(addr_parts) >= 2: self.addr_city_edit.setText(addr_parts[-1].strip())
        
        if o.get('city'): self.location_edit.setText(o.get('city'))
        
        parcels = o.get('parcels', [])
        nums = ', '.join(p['number'] if isinstance(p, dict) else str(p) for p in parcels)
        self.parcels_edit.setText(nums)
        
        street_dz_list = []
        if self.config.get('extract_parcel_address', True):
            for p_info in parcels:
                if isinstance(p_info, dict) and p_info.get('parcel_address'):
                    a = p_info.get('parcel_address')
                    if a and a not in street_dz_list: street_dz_list.append(a)
        
        self.street_dz_edit.setText(", ".join(street_dz_list))
        
        if parcels:
            bud_list, dem_list = [], []
            for p_info in parcels:
                num = p_info['number'] if isinstance(p_info, dict) else str(p_info)
                for known_p in self.parcels:
                    if known_p.get('number') == num:
                        c = known_p.get('category', '')
                        if 'Budowa' in c or 'Pe' in c or 'Przyłącze' in c or 'Przylacze' in c: bud_list.append(num)
                        if 'Demonta' in c or 'Demontaż' in c: dem_list.append(num)
            
            self.parcels_budowa_edit.setText(', '.join(set(bud_list)))
            self.parcels_demontaz_edit.setText(', '.join(set(dem_list)))
        self._auto_ownership()

    def _auto_ownership(self):
        sel = self.table_owners.selectedItems()
        if not sel: return
        
        row = sel[0].row()
        data = self.table_owners.item(row, 2).data(Qt.ItemDataRole.UserRole)
        if not data: return
        idx, specific_addr = data
        o = self.owners[idx]
        
        nums = [n.strip() for n in self.parcels_edit.text().split(',') if n.strip()]
        is_sole = o.get('share', '1/1') == '1/1'
        is_couple = o.get('is_couple', False)
        is_company_or_inst = o.get('is_company') or o.get('is_spolka') or o.get('is_institution') or o.get('is_church')
        plural = len(nums) > 1
        
        from utils.gender_utils import detect_gender
        g = detect_gender(o.get('first_name', ''))
        
        if is_company_or_inst: prefix = 'company'
        elif is_couple: prefix = 'couple'
        elif g == 'F': prefix = 'female'
        else: prefix = 'male'
        
        mid = 'sole' if is_sole else 'joint'
        suf = 'plur' if plural else 'sing'
        
        key = f"{prefix}_{mid}_{suf}"
        phrases = self.config.get('ownership_phrases', OwnershipPhraseDialog(self.config).default_phrases)
        phrase = phrases.get(key, "których są Państwo współwłaścicielami,")
        
        self.ownership_edit.setText(phrase)

    def _browse_template(self):
        from utils.templates import (
            EXAMPLES_FOLDER_NAMES,
            resolve_template_start_directory,
        )

        start_dir = resolve_template_start_directory(
            self.config,
            config_key='path_przyklady',
            folder_names=EXAMPLES_FOLDER_NAMES,
            current_path=self.template_edit.text(),
        )
        path, _ = QFileDialog.getOpenFileName(
            self,
            'Wybierz szablon',
            str(start_dir),
            'Word (*.docx)',
        )
        if path:
            self.template_edit.setText(path)

    def _set_default_template(self):
        from utils.templates import (
            EXAMPLES_FOLDER_NAMES,
            find_latest_file,
            resolve_template_start_directory,
        )

        przyk_path = resolve_template_start_directory(
            self.config,
            config_key='path_przyklady',
            folder_names=EXAMPLES_FOLDER_NAMES,
        )

        # Nie nadpisuj działającego wyboru dokonanego w tym module.
        cl_tmpl = self.template_edit.text().strip() or self.config.get('cover_letter_template', '')
        if not cl_tmpl or not Path(cl_tmpl).is_file():
            latest = find_latest_file(
                przyk_path,
                ["Pismo przewodnie", "pismo przewodnie"],
                (".docx",),
            )
            cl_tmpl = str(latest) if latest else ""
        self.template_edit.setText(cl_tmpl)

    def _get_params(self) -> dict:
        sender = self.config.get('sender', {})
        return {
            'date_str': self.date_edit.text().strip(),
            'place': self.sender_place_edit.text().strip(),
            'sender_name': sender.get('name', ''),
            'sender_street': sender.get('street', ''),
            'sender_city': sender.get('city', ''),
            'addressee_name': self.addr_name_edit.text().strip(),
            'addressee_street': self.addr_street_edit.text().strip(),
            'addressee_city': self.addr_city_edit.text().strip(),
            'location': self.location_edit.text().strip(),
            'street': self.street_dz_edit.text().strip(),
            'subject': self.subject_edit.text().strip(),
            'task_construction': self.task_budowa_edit.text().strip(),
            'task_demolition': self.task_demontaz_edit.text().strip(),
            'parcel_numbers_construction': [n.strip() for n in self.parcels_budowa_edit.text().split(',') if n.strip()],
            'parcel_numbers_demolition': [n.strip() for n in self.parcels_demontaz_edit.text().split(',') if n.strip()],
            'parcel_numbers': [n.strip() for n in self.parcels_edit.text().split(',') if n.strip()],
            'ownership_phrase': self.ownership_edit.text().strip(),
            'template_path': self.template_edit.text().strip(),
        }

    def _preview(self):
        # Podgląd ma prezentować dane źródłowe. Odmiana dotyczy wyłącznie
        # wartości podstawianych pod tagi w dokumencie DOCX.
        p = self._get_params()
        unique_parcels = list(dict.fromkeys(p['parcel_numbers']))
        all_nums_str = ', '.join(unique_parcels) or '—'
        parcel_type = 'działek nr' if len(unique_parcels) > 1 else 'działki nr'
        
        # Używamy tej samej funkcji formatującej (z limitem 34 znaków) w podglądzie
        try:
            from utils.docx_utils import format_long_name_for_cover
            formatted_name = format_long_name_for_cover(p['addressee_name'], max_len=34)
        except Exception:
            formatted_name = p['addressee_name']
        
        preview = (
            f"=== PISMO PRZEWODNIE ===\n"
            f"{p['place']}, {p['date_str']}\n\n"
            f"Sz. P.\n{formatted_name}\n{p['addressee_street']}\n{p['addressee_city']}\n\n"
            f"Zlokalizowanych w miejscowości {p['location']},\n"
            f"Ulica działki: {p['street']}\n"
            f"Teren {parcel_type} {all_nums_str}, {p['ownership_phrase']}."
        )
        self.preview_text.setText(preview)

    def _preview_docx(self):
        p = self._get_params()
        tmpl = p.get('template_path', '')
        if not tmpl or not Path(tmpl).exists():
            QMessageBox.warning(self, 'Brak szablonu', 'Wybierz lub załaduj domyślny szablon.')
            return
        self.preview_text.setText(read_docx_text(tmpl))

    def _generate(self):
        sel = self.table_owners.selectedItems()
        if sel:
            data = self.table_owners.item(sel[0].row(), 2).data(Qt.ItemDataRole.UserRole)
            if data:
                idx, specific_addr = data
                o = self.owners[idx]
                valid, reason = self._is_valid_for_gen(o, specific_addr)
                if not valid:
                    reply = QMessageBox.question(self, "Uwaga", f"Właściciel został pominięty przez filtr (Powód: {reason}).\nCzy chcesz WYMUSIĆ wygenerowanie pisma dla tej osoby?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
                    if reply == QMessageBox.StandardButton.No: return
        p = self._get_params()
        if not p['template_path'] or not Path(p['template_path']).is_file():
            self._set_default_template()
            p = self._get_params()
        if not p['template_path'] or not Path(p['template_path']).is_file():
            QMessageBox.warning(
                self,
                'Brak szablonu',
                'Wybierz poprawny plik szablonu Word dla pisma przewodniego.',
            )
            return

        out_path, _ = QFileDialog.getSaveFileName(self, 'Zapisz pismo', 'Pismo przewodnie.docx', 'Word (*.docx)')
        if not out_path: return
        
        from utils.docx_utils import generate_cover_letter
        try:
            ok = generate_cover_letter(
                template_path=p['template_path'], output_path=out_path, date_str=p['date_str'], place=p['place'],
                sender_name=p['sender_name'], sender_street=p['sender_street'], sender_city=p['sender_city'],
                addressee_salutation='Sz. P.', addressee_name=p['addressee_name'], addressee_street=p['addressee_street'],
                addressee_city=p['addressee_city'], location=p['location'],
                street=p['street'],
                subject=p['subject'],
                task_construction=p['task_construction'],
                task_demolition=p['task_demolition'],
                parcel_numbers_construction=p['parcel_numbers_construction'],
                parcel_numbers_demolition=p['parcel_numbers_demolition'],
                parcel_numbers=p['parcel_numbers'],
                ownership_phrase=p['ownership_phrase'], tag_map=self.config.get('cover_letter_tag_map'),
                unlock_docs=self.config.get("unlock_generated_docs", False),
                declension_options=self.config,
            )
            
            if ok: 
                if sel:
                    try:
                        data = self.table_owners.item(sel[0].row(), 2).data(Qt.ItemDataRole.UserRole)
                        idx, specific_addr = data
                        self._set_cover_done(self.owners[idx], specific_addr, True)
                        self.owners_changed.emit(self.owners)
                        self._refresh_owners_table()
                    except: pass
                QMessageBox.information(self, 'Gotowe', f'Plik zapisany.')
        except Exception as e:
            logger.exception("Błąd w trakcie ręcznego generowania Pisma Przewodniego.")
            QMessageBox.critical(self, 'Błąd Generowania', f'Wystąpił błąd podczas tworzenia pliku Word.\nSzczegóły zapisano w error_log.txt.\nBłąd: {e}')

    def _generate_selected(self):
        targets = []
        for r in range(self.table_owners.rowCount()):
            it = self.table_owners.item(r, 0)
            if it and it.checkState() == Qt.CheckState.Checked and not self.table_owners.isRowHidden(r):
                data = self.table_owners.item(r, 2).data(Qt.ItemDataRole.UserRole)
                if data:
                    idx, specific_addr = data
                    targets.append((self.owners[idx], specific_addr))
                    
        if not targets: return QMessageBox.warning(self, 'Brak', 'Nie zaznaczono nikogo ptaszkiem z widocznej listy.')
        self._generate_batch(targets)

    def _generate_all(self):
        if not self.owners:
            return QMessageBox.warning(
                self,
                'Brak właścicieli',
                'Wczytaj właścicieli na zakładce Wypisy.',
            )

        from utils.generation_targets import select_address_targets

        filter_text = self.search_owners_edit.text().strip().lower()

        def matches_filter(owner, specific_addr):
            fmt = self.config.get('couple_format_cover', 0)
            display_name = (
                owner.get('name_plural', owner.get('full_name', ''))
                if fmt == 0
                else owner.get('name_separate', owner.get('full_name', ''))
            )
            parcels_str = ' '.join(
                str(parcel.get('number', parcel))
                if isinstance(parcel, dict)
                else str(parcel)
                for parcel in owner.get('parcels', [])
            )
            search_target = f"{display_name} {specific_addr} {parcels_str}".lower()
            return filter_text in search_target

        targets = select_address_targets(
            self.owners,
            hide_done=self.chk_hide_generated.isChecked(),
            is_done=self._cover_done,
            matches_filter=matches_filter if filter_text else None,
        )
        if not targets:
            return QMessageBox.information(
                self,
                'Brak pism do wygenerowania',
                'Żaden adres nie spełnia obecnego filtra albo wszystkie pisma są już wygenerowane.',
            )
        self._generate_batch(targets)

    def _generate_batch(self, targets):
        if not targets:
            return QMessageBox.warning(
                self,
                'Brak',
                'Nie wybrano żadnego adresu do wygenerowania.',
            )

        p = self._get_params()
        if not p['template_path'] or not Path(p['template_path']).is_file():
            self._set_default_template()
            p = self._get_params()
        if not p['template_path'] or not Path(p['template_path']).is_file():
            return QMessageBox.warning(
                self,
                'Brak szablonu',
                'Wybierz poprawny plik szablonu Word dla pism przewodnich.',
            )

        out_dir = QFileDialog.getExistingDirectory(self, 'Folder wyjściowy')
        if not out_dir:
            return

        from utils.docx_utils import generate_cover_letter
        from utils.gender_utils import detect_gender
        success, errors = 0, []
        skipped_info = []
        phrases = self.config.get('ownership_phrases', OwnershipPhraseDialog(self.config).default_phrases)
        fmt = self.config.get('couple_format_cover', 0)
        for o, specific_addr in targets:
            valid, reason = self._is_valid_for_gen(o, specific_addr)
            if not valid: 
                skipped_info.append(f"{o.get('full_name', 'Nieznany')} ({reason})")
                continue
                
            name_line = o.get('name_plural', o.get('full_name', '')) if fmt == 0 else o.get('name_separate', o.get('full_name', ''))
            
            addr_parts = specific_addr.split(',')
            street_adr = addr_parts[0].strip() if addr_parts else ''
            city_adr = addr_parts[-1].strip() if len(addr_parts) > 1 else ''
            has_two_addresses = bool(o.get('address_2'))
            addr_suffix = ""
            if has_two_addresses:
                street_match = re.search(r'ul\.\s*([A-Za-zŻżÓóŁłĆćŃńŚśŹźŻż])', specific_addr, re.I)
                if street_match: addr_suffix = f" {street_match.group(1).upper()}"
                else: addr_suffix = f" 2" if specific_addr == o.get('address_2') else ""
            o_nums = [pr['number'] if isinstance(pr, dict) else str(pr) for pr in o.get('parcels', o.get('parcel_numbers', []))]
            
            is_sole = o.get('share', '1/1') == '1/1'
            is_couple = o.get('is_couple', False)
            is_company_or_inst = o.get('is_company') or o.get('is_spolka') or o.get('is_institution') or o.get('is_church')
            plural = len(o_nums) > 1
            g = detect_gender(o.get('first_name', ''))
            
            if is_company_or_inst: prefix = 'company'
            elif is_couple: prefix = 'couple'
            elif g == 'F': prefix = 'female'
            else: prefix = 'male'
            
            mid = 'sole' if is_sole else 'joint'
            suf = 'plur' if plural else 'sing'
            key = f"{prefix}_{mid}_{suf}"
            phrase = phrases.get(key, "których są Państwo współwłaścicielami,")
            
            short_name = self._get_short_name(o.get('first_name', ''), o.get('last_name', ''))
            fname = f"Pismo przewodnie {short_name}{addr_suffix}.docx".replace('  ', ' ')
            fname = re.sub(r'[<>:"/\\|?*]', '', fname)
            out_path = str(Path(out_dir) / fname)
            
            ulice_dz_list = []
            bud_list, dem_list = [], []
            if self.config.get('extract_parcel_address', True):
                for p_info in o.get('parcels', []):
                    if isinstance(p_info, dict):
                        num = p_info.get('number', '')
                        a = p_info.get('parcel_address')
                        if a and a not in ulice_dz_list: ulice_dz_list.append(a)
                        
                        for known_p in self.parcels:
                            if known_p.get('number') == num:
                                c = known_p.get('category', '')
                                if 'Budowa' in c or 'Pe' in c or 'Przyłącze' in c or 'Przylacze' in c: bud_list.append(num)
                                if 'Demonta' in c or 'Demontaż' in c: dem_list.append(num)
                                
            final_street_dz = ", ".join(ulice_dz_list) if ulice_dz_list else ""
            try:
                ok = generate_cover_letter(
                    template_path=p['template_path'], output_path=out_path, date_str=p['date_str'], place=p['place'],
                    sender_name=p['sender_name'], sender_street=p['sender_street'], sender_city=p['sender_city'],
                    addressee_salutation='Sz. P.', addressee_name=name_line, addressee_street=street_adr, addressee_city=city_adr,
                    location=o.get('city', '') or p['location'],
                    street=final_street_dz,
                    subject=p['subject'],
                    task_construction=p['task_construction'],
                    task_demolition=p['task_demolition'],
                    parcel_numbers_construction=list(set(bud_list)),
                    parcel_numbers_demolition=list(set(dem_list)),
                    parcel_numbers=o_nums,
                    ownership_phrase=phrase, tag_map=self.config.get('cover_letter_tag_map'),
                    unlock_docs=self.config.get("unlock_generated_docs", False),
                    declension_options=self.config,
                )
                if ok: 
                    success += 1
                    self._set_cover_done(o, specific_addr, True)
                else: 
                    errors.append(o.get('full_name', ''))
            except Exception as e:
                logger.exception(f"Błąd podczas batchowego generowania Pisma dla: {o.get('full_name', '')}")
                errors.append(f"{o.get('full_name', '')} [BŁĄD WORD]")
        self.owners_changed.emit(self.owners)
        self._refresh_owners_table()
        msg = f'Wygenerowano dokumentów: {success}.'
        if errors: 
            msg += f'\n\nBłędy generowania (Sprawdź logi w folderze "dane"):\n' + "\n".join(errors)
        if skipped_info:
            msg += f'\n\nPominięto przez filtr adresów/zmarłych (Ilość: {len(skipped_info)}):\n' + "\n".join(skipped_info[:15])
            if len(skipped_info) > 15: msg += "\n...i więcej."
            
        QMessageBox.information(self, 'Wynik', msg)
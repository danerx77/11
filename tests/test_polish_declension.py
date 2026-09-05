"""Testy automatycznej odmiany używanej wyłącznie przy podmianie tagów DOCX."""

from pathlib import Path
from tempfile import TemporaryDirectory
import unittest

try:
    from docx import Document
    from utils.docx_utils import apply_declension_preferences, generate_declaration
except ModuleNotFoundError:
    # Minimalne testy reguł nie wymagają python-docx. Testy integracyjne są
    # wykonywane automatycznie tam, gdzie zależność aplikacji jest dostępna.
    Document = None
    apply_declension_preferences = None
    generate_declaration = None

from utils.polish_declension import decline_city, decline_street


class CityDeclensionTests(unittest.TestCase):
    def test_regular_endings_cover_common_city_types(self):
        cases = {
            # -owo / -ewo, -ino / -yno, -ko / -no
            "Żukowo": "Żukowie",
            "Grabowo": "Grabowie",
            "Luzino": "Luzinie",
            "Chwaszczyno": "Chwaszczynie",
            "Kielno": "Kielnie",
            "Wicko": "Wicku",
            # liczba mnoga
            "Sierakowice": "Sierakowicach",
            "Kartuzy": "Kartuzach",
            "Czaple": "Czaplach",
            "Trąbki": "Trąbkach",
            # żeńskie
            "Stężyca": "Stężycy",
            "Kobylnica": "Kobylnicy",
            "Gdynia": "Gdyni",
            "Wieliczka": "Wieliczce",
            "Czerska": "Czerskiej",
            # męskie zakończone spółgłoską
            "Czersk": "Czersku",
            "Tczew": "Tczewie",
            "Przywidz": "Przywidzu",
            "Karsin": "Karsinie",
            "Żywiec": "Żywcu",
        }
        for source, expected in cases.items():
            with self.subTest(city=source):
                self.assertEqual(decline_city(source), expected)

    def test_regular_compound_and_hyphenated_city_names(self):
        cases = {
            "Stara Wieś": "Starej Wsi",
            "Tomaszów Mazowiecki": "Tomaszowie Mazowieckim",
            "Biała Podlaska": "Białej Podlaskiej",
            "Krynica-Zdrój": "Krynicy-Zdroju",
            "Skarżysko-Kamienna": "Skarżysku-Kamiennej",
            "Nowy Dwór Gdański": "Nowym Dworze Gdańskim",
            "Góra Kalwaria": "Górze Kalwarii",
            "Kostrzyn nad Odrą": "Kostrzynie nad Odrą",
            "Nowe Miasto nad Pilicą": "Nowym Mieście nad Pilicą",
        }
        for source, expected in cases.items():
            with self.subTest(city=source):
                self.assertEqual(decline_city(source), expected)

    def test_automatic_dictionary_precedes_county_check_and_preserves_case(self):
        self.assertEqual(
            decline_city("Pruszcz Gdański"), "Pruszczu Gdańskim"
        )
        self.assertEqual(decline_city("80-000 Stara Wieś"), "80-000 Starej Wsi")
        self.assertEqual(decline_city("Stalowa Wola"), "Stalowej Woli")
        self.assertEqual(decline_city("Kędzierzyn-Koźle"), "Kędzierzynie-Koźlu")
        self.assertEqual(
            decline_city("Wysokie Mazowieckie"), "Wysokiem Mazowieckiem"
        )
        self.assertEqual(decline_city("Kwidzyn"), "Kwidzynie")
        self.assertEqual(decline_city("RUDA ŚLĄSKA"), "RUDZIE ŚLĄSKIEJ")
        self.assertEqual(decline_city("ŻUKOWO"), "ŻUKOWIE")
        self.assertEqual(decline_city("gdański"), "gdański")


class StreetDeclensionTests(unittest.TestCase):
    def test_regular_street_names_and_multiword_names(self):
        cases = {
            "Spokojna": "Spokojnej",
            "Polna": "Polnej",
            "Długa": "Długiej",
            "Szeroka": "Szerokiej",
            "Gdańska": "Gdańskiej",
            "ul. Spokojna": "ul. Spokojnej",
            "ulica Miła": "ulica Miłej",
            "UL. SPOKOJNA": "UL. SPOKOJNEJ",
            "Nowa Wieś": "Nowej Wsi",
            "Stara Droga": "Starej Drodze",
            "Plac Wolności": "Placu Wolności",
            "Aleja Róż": "Alei Róż",
            "Słowacka-Krasickiego": "Słowackiej-Krasickiego",
            "ul. Spokojna 12A/3": "ul. Spokojnej 12A/3",
        }
        for source, expected in cases.items():
            with self.subTest(street=source):
                self.assertEqual(decline_street(source), expected)

    def test_person_date_and_already_declined_names_are_left_intact(self):
        cases = (
            "Mickiewicza",
            "Andersa",
            "Marcina Kasprzaka",
            "Generała Andersa",
            "3 Maja",
            "11 Listopada",
            "Świętego Jana",
            "Jana Pawła II",
            "Marii Skłodowskiej-Curie",
            "Księdza Popiełuszki",
            "Piłsudskiego",
            "Spokojnej",
        )
        for source in cases:
            with self.subTest(street=source):
                self.assertEqual(decline_street(source), source)

    def test_address_suffix_after_comma_is_not_lost_or_changed(self):
        source = "ul. Spokojna, 80-000 Gdynia"
        self.assertEqual(
            decline_street(source), "ul. Spokojnej, 80-000 Gdynia"
        )

    @unittest.skipUnless(Document, "python-docx is not installed")
    def test_preferences_only_transform_tag_values(self):
        location = "Nowa Wieś"
        street = "ul. Testowa"
        preferences = {
            "decl_location_locative": True,
            "decl_decline_streets": True,
        }

        tag_location, tag_street, tag_county = apply_declension_preferences(
            location, street, "", preferences
        )

        self.assertEqual(tag_location, "Nowej Wsi")
        self.assertEqual(tag_street, "ul. Testowej")
        self.assertEqual(tag_county, "")
        self.assertEqual(location, "Nowa Wieś")
        self.assertEqual(street, "ul. Testowa")

    @unittest.skipUnless(Document, "python-docx is not installed")
    def test_docx_keeps_raw_address_and_declines_only_tags(self):
        with TemporaryDirectory() as temp_dir:
            template = Path(temp_dir) / "template.docx"
            output = Path(temp_dir) / "output.docx"
            document = Document()
            document.add_paragraph("Miejscowość: <Miejscowość działki:>")
            document.add_paragraph("Ulica: <Ulica>")
            document.add_paragraph("Adres: <Adres>")
            document.save(template)

            source_location = "Nowa Wieś"
            source_street = "ul. Testowa"
            success = generate_declaration(
                str(template),
                str(output),
                location=source_location,
                street=source_street,
                declension_options={
                    "decl_location_locative": True,
                    "decl_decline_streets": True,
                },
            )

            self.assertTrue(success)
            self.assertEqual(source_location, "Nowa Wieś")
            self.assertEqual(source_street, "ul. Testowa")
            output_text = "\n".join(
                paragraph.text for paragraph in Document(output).paragraphs
            )
            self.assertIn("Miejscowość: Nowej Wsi", output_text)
            self.assertIn("Ulica: ul. Testowej", output_text)
            self.assertIn("Adres: ul. Testowa, Nowa Wieś", output_text)


if __name__ == "__main__":
    unittest.main()

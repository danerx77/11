"""
parcel_list.py – Zakładka lista działek   Lista Działek
"""
import json
from pathlib import Path

from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QLineEdit,
    QTableWidget, QTableWidgetItem, QFileDialog, QMessageBox, QComboBox,
    QGroupBox, QSplitter, QTextEdit, QDialog, QDialogButtonBox, QFormLayout,
    QHeaderView, QAbstractItemView
)
from PySide6.QtCore import Qt, Signal
from PySide6.QtGui import QColor, QFont, QShortcut, QKeySequence

from utils.parcel_sorting import parcel_sort_key

CATEGORY_COLORS = {
    'Demontaż': QColor('#ff6b35'),
    'Budowa': QColor('#00b4d8'),
    'Przyłącze': QColor('#2ecc71'),
    'Budowa, Demontaż': QColor('#9b5de5'),
    'Budowa, Przyłącze': QColor('#1abc9c'),
    'Demontaż, Przyłącze': QColor('#f1c40f'),
    'Budowa, Demontaż, Przyłącze': QColor('#e84393'),
}

class AddParcelDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle('Dodaj działkę (Ręcznie)')
        self.setMinimumWidth(400)
        self.setStyleSheet("""
            QGroupBox { font-weight: bold; font-size: 14px; margin-top: 15px; border: 1px solid palette(mid); border-radius: 6px; }
            QGroupBox::title { subcontrol-origin: margin; left: 10px; padding: 0 5px 0 5px; }
            QLineEdit { padding: 6px; font-size: 13px; border-radius: 4px; border: 1px solid palette(mid); }
            QCheckBox { font-size: 14px; padding: 2px; }
            QPushButton#btn_add { background-color: #2ecc71; color: white; font-weight: bold; padding: 8px; border-radius: 4px; border: none; }
            QPushButton#btn_add:hover { background-color: #27ae60; }
            QPushButton#btn_cancel { background-color: #e74c3c; color: white; font-weight: bold; padding: 8px; border-radius: 4px; border: none; }
            QPushButton#btn_cancel:hover { background-color: #c0392b; }
        """)
        
        from PySide6.QtWidgets import QVBoxLayout, QHBoxLayout, QCheckBox, QGroupBox, QFormLayout
        main_layout = QVBoxLayout(self)
        main_layout.setSpacing(15)

        group_info = QGroupBox("Podstawowe informacje")
        layout_info = QFormLayout(group_info)
        layout_info.setSpacing(10)

        self.number_edit = QLineEdit()
        self.number_edit.setPlaceholderText('np. 453/5')
        layout_info.addRow('Numer działki:', self.number_edit)

        self.precinct_edit = QLineEdit()
        self.precinct_edit.setPlaceholderText('np. Polki')
        layout_info.addRow('Obręb:', self.precinct_edit)
        
        main_layout.addWidget(group_info)

        group_cat = QGroupBox("Kategorie (Możesz zaznaczyć wiele)")
        layout_cat = QHBoxLayout(group_cat)
        
        self.chk_budowa = QCheckBox('Budowa')
        self.chk_demontaz = QCheckBox('Demontaż')
        self.chk_przylacze = QCheckBox('Przyłącze')
        
        layout_cat.addWidget(self.chk_budowa)
        layout_cat.addWidget(self.chk_demontaz)
        layout_cat.addWidget(self.chk_przylacze)
        
        main_layout.addWidget(group_cat)

        btn_layout = QHBoxLayout()
        btn_layout.addStretch()
        
        btn_cancel = QPushButton("Anuluj")
        btn_cancel.setObjectName("btn_cancel")
        btn_cancel.clicked.connect(self.reject)
        btn_layout.addWidget(btn_cancel)
        
        btn_ok = QPushButton("Dodaj Działkę")
        btn_ok.setObjectName("btn_add")
        btn_ok.clicked.connect(self.accept)
        btn_layout.addWidget(btn_ok)
        
        main_layout.addLayout(btn_layout)

    def get_values(self):
        cats = []
        if self.chk_budowa.isChecked(): cats.append('Budowa')
        if self.chk_demontaz.isChecked(): cats.append('Demontaż')
        if self.chk_przylacze.isChecked(): cats.append('Przyłącze')
        
        cat = ', '.join(cats)
        if not cat: cat = 'Budowa'
        
        return [{'number': self.number_edit.text().strip(), 'precinct': self.precinct_edit.text().strip(), 'category': cat}]


class ParcelListWidget(QWidget):
    parcels_changed = Signal(list)

    def __init__(self, config: dict, parent=None):
        super().__init__(parent)
        self.config = config
        self.parcels = [] 
        self.setAcceptDrops(True)
        self._build_ui()

    def dragEnterEvent(self, e):
        if e.mimeData().hasUrls(): 
            e.acceptProposedAction()

    def dropEvent(self, e):
        for url in e.mimeData().urls():
            if url.isLocalFile():
                path = url.toLocalFile()
                ext = Path(path).suffix.lower()
                if ext in ['.txt', '.docx', '.xlsx', '.xlsm', '.xls']:
                    self.path_edit.setText(path)
                    self._load_file(path)
                break

    def _build_ui(self):
        main_layout = QHBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        splitter = QSplitter(Qt.Orientation.Horizontal)
        main_layout.addWidget(splitter)

        left = QWidget()
        left_layout = QVBoxLayout(left)
        left_layout.setContentsMargins(8, 8, 4, 8)

        header_row = QHBoxLayout()
        header_label = QLabel('📋 Lista działek')
        header_label.setStyleSheet('font-size:15px; font-weight:700;')
        header_row.addWidget(header_label)
        header_row.addStretch()

        self.search_edit = QLineEdit()
        self.search_edit.setPlaceholderText('Szukaj działki, np. 123/1...')
        self.search_edit.setMinimumWidth(180)
        self.search_edit.textChanged.connect(self._apply_filter)
        header_row.addWidget(QLabel('Szukaj:'))
        header_row.addWidget(self.search_edit)

        self.filter_combo = QComboBox()
        self.filter_combo.addItems(['Wszystkie', 'Demontaż', 'Budowa', 'Przyłącze'])
        saved_filter = self.config.get('parcel_list_filter', 'Wszystkie')
        if saved_filter in [self.filter_combo.itemText(i) for i in range(self.filter_combo.count())]:
            self.filter_combo.setCurrentText(saved_filter)

        self.sort_combo = QComboBox()
        self.sort_combo.addItems(['Domyślne', 'Rosnąco', 'Malejąco'])
        saved_sort = self.config.get('parcel_list_sort', 'Domyślne')
        if saved_sort in [self.sort_combo.itemText(i) for i in range(self.sort_combo.count())]:
            self.sort_combo.setCurrentText(saved_sort)

        self.filter_combo.currentTextChanged.connect(self._remember_list_preferences)
        self.sort_combo.currentTextChanged.connect(self._remember_list_preferences)
        
        header_row.addWidget(QLabel('Filtr:'))
        header_row.addWidget(self.filter_combo)
        header_row.addWidget(QLabel(' Sortowanie:'))
        header_row.addWidget(self.sort_combo)
        left_layout.addLayout(header_row)

        self.table = QTableWidget(0, 4)
        self.table.setHorizontalHeaderLabels(['Nr działki', 'Obręb', 'Kategoria', ''])
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeMode.ResizeToContents)
        self.table.setAlternatingRowColors(True)
        self.table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self.table.setEditTriggers(QAbstractItemView.EditTrigger.DoubleClicked)

        QShortcut(QKeySequence("Delete"), self.table).activated.connect(self._delete_selected)

        self.table.horizontalHeader().setSectionsMovable(True)
        table_state_parcels_hex = self.config.get('table_state_parcels', '')
        if table_state_parcels_hex:
            from PySide6.QtCore import QByteArray
            self.table.horizontalHeader().restoreState(QByteArray.fromHex(table_state_parcels_hex.encode()))
        self.table.horizontalHeader().sectionResized.connect(lambda *args: self.config.update({'table_state_parcels': self.table.horizontalHeader().saveState().toHex().data().decode()}))
        self.table.horizontalHeader().sectionMoved.connect(lambda *args: self.config.update({'table_state_parcels': self.table.horizontalHeader().saveState().toHex().data().decode()}))
        left_layout.addWidget(self.table)

        btn_row = QHBoxLayout()
        self.btn_add = QPushButton('+ Dodaj')
        self.btn_add.setObjectName('btn_primary')
        self.btn_add.clicked.connect(self._add_parcel)
        btn_row.addWidget(self.btn_add)

        self.btn_delete = QPushButton('🗑 Usuń zaznaczone (Delete)')
        self.btn_delete.setObjectName('btn_danger')
        self.btn_delete.clicked.connect(self._delete_selected)
        btn_row.addWidget(self.btn_delete)

        self.btn_clear = QPushButton('🗑 Wyczyść wszystko')
        self.btn_clear.clicked.connect(self._clear_all)
        btn_row.addWidget(self.btn_clear)

        btn_row.addStretch()
        self.lbl_count = QLabel('Działek: 0')
        self.lbl_count.setStyleSheet('color:#aaa;')
        btn_row.addWidget(self.lbl_count)
        left_layout.addLayout(btn_row)

        splitter.addWidget(left)

        right = QWidget()
        right_layout = QVBoxLayout(right)
        right_layout.setContentsMargins(4, 8, 8, 8)

        import_box = QGroupBox('Import z pliku')
        import_layout = QVBoxLayout(import_box)

        self.path_edit = QLineEdit()
        self.path_edit.setReadOnly(True)
        self.path_edit.setPlaceholderText('Wybierz plik lub przeciągnij tutaj...')
        import_layout.addWidget(self.path_edit)

        btn_browse = QPushButton('📂 Wybierz plik (txt/docx/xlsx)')
        btn_browse.clicked.connect(self._browse_file)
        import_layout.addWidget(btn_browse)

        btn_load = QPushButton('⬆ Wczytaj z folderu lista/ projektu')
        btn_load.clicked.connect(self._load_from_project)
        import_layout.addWidget(btn_load)

        right_layout.addWidget(import_box)

        paste_box = QGroupBox('Wklej / wpisz listę')
        paste_layout = QVBoxLayout(paste_box)
        lbl_hint = QLabel('Format: wpisz "Demontaż", "Budowa" lub "Przyłącze" jako nagłówek,\na pod nim numery działek.')
        lbl_hint.setWordWrap(True)
        lbl_hint.setStyleSheet('color:#aaa; font-size:11px;')
        paste_layout.addWidget(lbl_hint)

        self.paste_edit = QTextEdit()
        self.paste_edit.setPlaceholderText('Demontaż\n453/5\n\nBudowa\n676/3\n\nPrzyłącze\n100/1')
        self.paste_edit.setMaximumHeight(180)
        paste_layout.addWidget(self.paste_edit)

        btn_parse = QPushButton('✅ Przetwórz i dodaj do listy')
        btn_parse.setObjectName('btn_primary')
        btn_parse.clicked.connect(self._parse_paste)
        paste_layout.addWidget(btn_parse)

        right_layout.addWidget(paste_box)

        export_box = QGroupBox('Eksport')
        export_layout = QVBoxLayout(export_box)

        btn_export_docx = QPushButton('💾 Eksportuj do DOCX')
        btn_export_docx.clicked.connect(self._export_docx)
        export_layout.addWidget(btn_export_docx)

        btn_export_txt = QPushButton('💾 Eksportuj do TXT')
        btn_export_txt.clicked.connect(self._export_txt)
        export_layout.addWidget(btn_export_txt)

        right_layout.addWidget(export_box)
        right_layout.addStretch()

        splitter.addWidget(right)
        splitter.setSizes([500, 300])

    def _refresh_table(self):
        filter_cat = self.filter_combo.currentText()
        sort_mode = self.sort_combo.currentText()
        search_text = self.search_edit.text().strip().lower() if hasattr(self, 'search_edit') else ''
        
        self.table.setRowCount(0)
        
        display_list = []
        for p in self.parcels:
            if filter_cat != 'Wszystkie' and filter_cat not in p['category']:
                continue
            if search_text:
                target = f"{p.get('number','')} {p.get('precinct','')} {p.get('category','')}".lower()
                if search_text not in target:
                    continue
            display_list.append(p)

        if sort_mode != 'Domyślne':
            display_list.sort(
                key=lambda parcel: parcel_sort_key(parcel.get('number', '')),
                reverse=(sort_mode == 'Malejąco'),
            )

        shown = 0
        for p in display_list:
            row = self.table.rowCount()
            self.table.insertRow(row)

            num_item = QTableWidgetItem(p['number'])
            prec_item = QTableWidgetItem(p.get('precinct', ''))
            
            # Wstawienie ComboBoxa zamiast samego tekstu
            cat_combo = QComboBox()
            cat_combo.addItems([
                'Budowa', 'Demontaż', 'Przyłącze', 
                'Budowa, Demontaż', 'Budowa, Przyłącze', 
                'Demontaż, Przyłącze', 'Budowa, Demontaż, Przyłącze'
            ])
            cat_combo.setCurrentText(p['category'])
            
            color = CATEGORY_COLORS.get(p['category'], QColor('#555'))
            cat_combo.setStyleSheet(f"QComboBox {{ color: {color.name()}; font-weight: bold; }} QComboBox QAbstractItemView {{ selection-background-color: #34495e; }}")
            
            idx = self.parcels.index(p)
            cat_combo.currentTextChanged.connect(lambda text, i=idx, cb=cat_combo: self._change_category(i, text, cb))

            btn_del = QPushButton('🗑️')
            btn_del.setFixedWidth(40)
            btn_del.setStyleSheet('color: #d63031; font-size: 14px; font-weight: bold; padding: 2px;')
            btn_del.clicked.connect(lambda _, i=idx: self._delete_by_index(i))
            
            num_item.setData(Qt.ItemDataRole.UserRole, idx)

            self.table.setItem(row, 0, num_item)
            self.table.setItem(row, 1, prec_item)
            self.table.setCellWidget(row, 2, cat_combo)
            self.table.setCellWidget(row, 3, btn_del)
            shown += 1

        self.lbl_count.setText(f'Działek: {len(self.parcels)} (widocznych: {shown})')
        self.parcels_changed.emit(self.parcels)

    def _change_category(self, idx, new_cat, combo):
        if 0 <= idx < len(self.parcels):
            self.parcels[idx]['category'] = new_cat
            color = CATEGORY_COLORS.get(new_cat, QColor('#555'))
            combo.setStyleSheet(f"QComboBox {{ color: {color.name()}; font-weight: bold; }} QComboBox QAbstractItemView {{ selection-background-color: #34495e; }}")
            self._save_to_project_state()
            self.parcels_changed.emit(self.parcels)

    def _remember_list_preferences(self, *_args):
        """Przechowuje aktywny filtr i sposób sortowania między uruchomieniami."""
        self.config['parcel_list_filter'] = self.filter_combo.currentText()
        self.config['parcel_list_sort'] = self.sort_combo.currentText()
        self._refresh_table()

    def _apply_filter(self):
        self._refresh_table()

    def _add_or_update_parcel(self, num: str, cat: str):
        cats_to_add = [x.strip() for x in cat.replace(' i ', ',').split(',')]
        expanded_cats = set()
        for c in cats_to_add:
            if c in ['Pełna', 'Oba', 'Budowa i Demontaż', 'Budowa, Demontaż']:
                expanded_cats.update(['Budowa', 'Demontaż'])
            else:
                expanded_cats.add(c)

        for p in self.parcels:
            if p['number'] == num:
                current_cats = [x.strip() for x in p['category'].replace(' i ', ',').split(',')]
                valid_cats = set()
                for c in current_cats:
                    if 'Budowa' in c: valid_cats.add('Budowa')
                    if 'Demontaż' in c: valid_cats.add('Demontaż')
                    if 'Przyłącze' in c: valid_cats.add('Przyłącze')
                    if 'Pełna' in c: valid_cats.update(['Budowa', 'Demontaż'])

                valid_cats.update(expanded_cats)
                order = ['Budowa', 'Demontaż', 'Przyłącze']
                final_cats = [x for x in order if x in valid_cats]
                p['category'] = ', '.join(final_cats)
                return
        
        order = ['Budowa', 'Demontaż', 'Przyłącze']
        final_cats = [x for x in order if x in expanded_cats]
        if not final_cats: final_cats = ['Budowa']
        self.parcels.append({'number': num, 'precinct': '', 'category': ', '.join(final_cats)})

    def _add_parcel(self):
        dlg = AddParcelDialog(self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            vals = dlg.get_values()
            for v in vals:
                if v['number']: self._add_or_update_parcel(v['number'], v['category'])
            self._refresh_table()
            self._save_to_project_state()

    def _delete_selected(self):
        indices_to_delete = set()
        for item in self.table.selectedItems():
            idx_data = self.table.item(item.row(), 0).data(Qt.ItemDataRole.UserRole)
            if idx_data is not None:
                indices_to_delete.add(idx_data)
        
        for idx in sorted(list(indices_to_delete), reverse=True):
            if idx < len(self.parcels):
                self.parcels.pop(idx)
                
        self._refresh_table()
        self._save_to_project_state()

    def _delete_by_index(self, idx: int):
        if 0 <= idx < len(self.parcels):
            self.parcels.pop(idx)
            self._refresh_table()
            self._save_to_project_state()

    def _clear_all(self):
        if QMessageBox.question(self, 'Wyczyść', 'Usunąć całą listę działek?', QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No) == QMessageBox.StandardButton.Yes:
            self.parcels.clear()
            self._refresh_table()
            self._save_to_project_state()

    def _browse_file(self):
        path, _ = QFileDialog.getOpenFileName(self, 'Otwórz listę działek', '', 'Pliki tekstowe (*.txt *.docx *.xlsx *.xlsm *.xls);;Wszystkie (*.*)')
        if path:
            self.path_edit.setText(path)
            self._load_file(path)

    def _load_from_project(self):
        last_path = self.config.get('last_project_path', '')
        if not last_path: return QMessageBox.warning(self, 'Brak projektu', 'Nie wybrano aktywnego projektu.')
        lista_dir = Path(last_path) / 'lista'
        if not lista_dir.exists(): return QMessageBox.warning(self, 'Brak folderu', f'Brak folderu lista/ w projekcie.')
        files = [f for f in lista_dir.glob('*') if f.suffix.lower() in ('.txt', '.docx', '.xlsx', '.xlsm', '.xls')]
        if not files: return QMessageBox.information(self, 'Brak plików', f'Brak plików w {lista_dir}')
        if len(files) == 1: self._load_file(str(files[0]))
        else:
            from PySide6.QtWidgets import QInputDialog
            choice, ok = QInputDialog.getItem(self, 'Wybierz plik', 'Pliki w folderze lista/:', [f.name for f in files], 0, False)
            if ok:
                idx = [f.name for f in files].index(choice)
                self._load_file(str(files[idx]))

    def _load_file(self, filepath: str):
        try:
            from utils.pdf_utils import parse_parcel_list_file
            result = parse_parcel_list_file(filepath)
            for num in result['demolition']: self._add_or_update_parcel(num, 'Demontaż')
            for num in result['construction']: self._add_or_update_parcel(num, 'Budowa')
            for num in result['connection']: self._add_or_update_parcel(num, 'Przyłącze')
            for num in result['full']: self._add_or_update_parcel(num, 'Budowa, Demontaż')
            self._refresh_table()
            self._save_to_project_state()
            QMessageBox.information(self, 'Wczytano', f'Pomyślnie zaktualizowano listę z pliku.')
        except Exception as e:
            QMessageBox.critical(self, 'Błąd', f'Błąd wczytywania pliku:\n{e}')

    def _parse_paste(self):
        text = self.paste_edit.toPlainText()
        if not text.strip(): return
        from utils.pdf_utils import parse_parcel_list_text
        result = parse_parcel_list_text(text)
        for num in result['demolition']: self._add_or_update_parcel(num, 'Demontaż')
        for num in result['construction']: self._add_or_update_parcel(num, 'Budowa')
        for num in result['connection']: self._add_or_update_parcel(num, 'Przyłącze')
        for num in result['full']: self._add_or_update_parcel(num, 'Budowa, Demontaż')
        self._refresh_table()
        self._save_to_project_state()
        self.paste_edit.clear()
        QMessageBox.information(self, 'Dodano', f'Zaktualizowano listę działek.')

    def _export_docx(self):
        path, _ = QFileDialog.getSaveFileName(self, 'Zapisz listę działek', 'lista_dzialek.docx', 'Word (*.docx)')
        if not path: return
        try:
            import docx as python_docx
            doc = python_docx.Document()
            doc.add_heading('Lista działek', 0)
            for cat in ['Demontaż', 'Budowa', 'Przyłącze']:
                cat_parcels = [p for p in self.parcels if cat in p['category']]
                if cat_parcels:
                    doc.add_heading(cat, level=1)
                    seen = set()
                    for p in cat_parcels:
                        if p['number'] in seen: continue
                        seen.add(p['number'])
                        line = p['number']
                        if p.get('precinct'): line += f" (obręb: {p['precinct']})"
                        doc.add_paragraph(line, style='List Bullet')
            doc.save(path)
            QMessageBox.information(self, 'Zapisano', f'Plik zapisany:\n{path}')
        except Exception as e:
            QMessageBox.critical(self, 'Błąd', str(e))

    def _export_txt(self):
        path, _ = QFileDialog.getSaveFileName(self, 'Zapisz listę działek', 'lista_dzialek.txt', 'Tekst (*.txt)')
        if not path: return
        lines = []
        for cat in ['Demontaż', 'Budowa', 'Przyłącze']:
            cat_parcels = [p for p in self.parcels if cat in p['category']]
            if cat_parcels:
                lines.append(cat)
                seen = set()
                for p in cat_parcels:
                    if p['number'] in seen: continue
                    seen.add(p['number'])
                    lines.append(p['number'])
                lines.append('')
        Path(path).write_text('\n'.join(lines), encoding='utf-8')
        QMessageBox.information(self, 'Zapisano', f'Plik zapisany:\n{path}')

    def set_project(self, project: dict):
        self._load_from_project_state(project.get('path', ''))

    def _save_to_project_state(self):
        last_path = self.config.get('last_project_path', '')
        if not last_path: return
        state_file = Path(last_path) / 'parcels_state.json'
        try:
            with open(state_file, 'w', encoding='utf-8') as f: json.dump(self.parcels, f, ensure_ascii=False, indent=2)
        except: pass

    def _load_from_project_state(self, project_path: str):
        self.parcels.clear()
        if not project_path:
            self._refresh_table()
            return
        state_file = Path(project_path) / 'parcels_state.json'
        if state_file.exists():
            try:
                with open(state_file, 'r', encoding='utf-8') as f: self.parcels = json.load(f)
            except: pass
        self._refresh_table()

    def get_parcels(self) -> list:
        return self.parcels